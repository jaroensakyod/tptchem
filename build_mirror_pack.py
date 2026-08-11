#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""CurioNest Mirror Pack builder — schema ใหม่ (ล้อ OpenStax) → docx → PDF.

Schema (mirror-json/*.json):
  sections[]:
    header            → cover header
    learning_objectives → objectives block
    concept           → concept banner + body paras + worked_example + common_mistakes
    practice          → practice items (short/mc/calc) + answers
    vocabulary        → key terms table
    answer_key        → answer key section (จาก answers ที่ฝังใน items)

Reuses omml.py ($...$ OMML) + font vertAlign (HTML sub/sup) — same proven
equation pipeline. US Letter, Georgia, page border, header/footer.

Usage: python build_mirror_pack.py mirror-json/b01-mirror.json
"""
import json, os, re, subprocess, sys

BASE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INDIGO = RGBColor(0x31, 0x2E, 0x81)
PURPLE = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x33, 0x41, 0x55)
GRAY = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0x9F, 0x12, 0x39)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE = RGBColor(0xC2, 0x41, 0x0C)

# ---------------------------------------------------------------- helpers
def add_chem_text(par, text, size=10.5, bold=False, color=None, italic=False):
    """$latex$ → OMML display; <sub>/<sup>/<b>/<i>/<br> → font vertAlign."""
    from lxml import etree
    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    for i, seg in enumerate(re.split(r"(\$[^$]+\$)", text)):
        if not seg:
            continue
        if seg.startswith("$") and seg.endswith("$") and len(seg) > 2:
            latex = seg[1:-1]
            try:
                from omml import latex_to_omml
                body = latex_to_omml(latex)
                omath = etree.fromstring(
                    f'<m:oMath xmlns:m="{M_NS}">'
                    f'<m:r><m:rPr><m:sz m:val="{int(size * 2)}"/></m:rPr></m:r>{body}'
                    f"</m:oMath>")
                par._p.append(omath)
            except Exception:
                run = par.add_run(seg)
                run.font.size = Pt(size)
                run.font.bold = bold
                if color:
                    run.font.color.rgb = color
            continue
        _add_html_runs(par, seg, size, bold, color, italic)
    return par

def _add_html_runs(par, text, size=10.5, bold=False, color=None, italic=False):
    stack = []
    pos = 0
    for m in re.finditer(r"<(/?)(sub|sup|b|i|br)([^>]*)>", text):
        if m.start() > pos:
            run = par.add_run(text[pos:m.start()])
            run.font.size = Pt(size)
            run.font.bold = bold or "b" in stack
            run.font.italic = italic or "i" in stack
            if color:
                run.font.color.rgb = color
            if "sub" in stack:
                run.font.subscript = True
            if "sup" in stack:
                run.font.superscript = True
        closing, tag, _ = m.groups()
        if tag == "br":
            par.add_run().add_break()
        elif closing:
            if tag in stack:
                stack.remove(tag)
        else:
            stack.append(tag)
        pos = m.end()
    if pos < len(text):
        run = par.add_run(text[pos:])
        run.font.size = Pt(size)
        run.font.bold = bold or "b" in stack
        run.font.italic = italic or "i" in stack
        if color:
            run.font.color.rgb = color
        if "sub" in stack:
            run.font.subscript = True
        if "sup" in stack:
            run.font.superscript = True

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcPr.append(shd)

def cell_borders(cell, color="999999", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.65)
    sec.top_margin, sec.bottom_margin = Inches(0.55), Inches(0.6)
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(2)
    return doc

def add_page_border(doc, color="4F46E5", sz="10"):
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), color)
        pgBorders.append(el)
    sectPr.append(pgBorders)

def add_header_footer(doc, brand_text):
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(brand_text)
    r.font.size = Pt(7.5)
    r.font.color.rgb = GRAY
    r.font.name = "Georgia"
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = fp.add_run("© 2026 CurioNest · For classroom use only · Page ")
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = GRAY
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    fp.add_run()._r.append(fld1)
    fp.add_run()._r.append(instr)
    fp.add_run()._r.append(fld2)

def para(doc, text, size=10, color=SLATE, bold=False, italic=False, after=3,
         align=WD_ALIGN_PARAGRAPH.LEFT, indent=None, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if before:
        p.paragraph_format.space_before = Pt(before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    add_chem_text(p, text, size=size, bold=bold, color=color, italic=italic)
    return p

def heading(doc, text, size=12, color=INDIGO, bold=True, after=4, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    add_chem_text(p, text, size=size, bold=bold, color=color)
    return p

# ---------------------------------------------------------------- section renderers
def render_header(doc, s):
    para(doc, "", size=4, after=0)
    para(doc, "CURIONEST", size=13, color=INDIGO, bold=True, after=4,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, s.get("title", ""), size=19, color=INDIGO, bold=True, after=3,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, s.get("subtitle", ""), size=10, color=GRAY, italic=True, after=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    # What's inside table
    items = [
        "✓  Complete concept notes with worked examples",
        "✓  Practice sets in varied formats (short answer, MCQ, calculations)",
        "✓  Key terms glossary + full answer key with worked solutions",
        "✓  Original figures and Word-native equations — print-ready & editable"
    ]
    tbl = doc.add_table(rows=len(items) + 1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade_cell(tbl.cell(0, 0), "EEF2FF")
    hp = tbl.cell(0, 0).paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("WHAT'S INSIDE")
    hr.font.bold = True; hr.font.size = Pt(10); hr.font.color.rgb = INDIGO
    for i, item in enumerate(items):
        cp_ = tbl.cell(i + 1, 0).paragraphs[0]
        r = cp_.add_run(item)
        r.font.size = Pt(9); r.font.color.rgb = SLATE
    para(doc, "Made by a chemistry teacher — aligned to the standard US chemistry sequence.",
         size=8.5, color=GRAY, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

def render_learning_objectives(doc, s):
    heading(doc, s.get("title", "Learning Objectives"), size=11.5)
    for i, obj in enumerate(s.get("items", [])):
        para(doc, f"{i+1}.  {obj}", size=9.8, indent=0.25, after=1)

def render_concept(doc, s):
    # banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, "4F46E5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(s.get("title", ""))
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = WHITE
    r.font.name = "Georgia"
    # body
    for b in s.get("body", []):
        para(doc, b, size=9.8, after=4)
    we = s.get("worked_example")
    if we:
        tbl2 = doc.add_table(rows=1, cols=1)
        tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
        c2 = tbl2.cell(0, 0)
        shade_cell(c2, "FFF7ED")
        cell_borders(c2, color="FDBA74", sz="6")
        p2 = c2.paragraphs[0]
        add_chem_text(p2, "★  Worked Example", size=9.5, bold=True, color=ORANGE)
        p2b = c2.add_paragraph()
        add_chem_text(p2b, we.get("question", ""), size=9.3, color=SLATE)
        p2c = c2.add_paragraph()
        r2c = p2c.add_run("Solution: ")
        r2c.font.bold = True; r2c.font.size = Pt(9.3); r2c.font.color.rgb = ORANGE
        add_chem_text(p2c, we.get("solution", ""), size=9.3, color=SLATE)
    for cm in s.get("common_mistakes", []):
        tbl3 = doc.add_table(rows=1, cols=1)
        tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
        c3 = tbl3.cell(0, 0)
        shade_cell(c3, "FEF2F2")
        cell_borders(c3, color="B91C1C", sz="6")
        p3 = c3.paragraphs[0]
        r3 = p3.add_run("★  Common mistake: ")
        r3.font.bold = True; r3.font.size = Pt(9.5); r3.font.color.rgb = RED
        add_chem_text(p3, cm, size=9.5, color=SLATE)


def add_table_borders(tbl):
    """Give every cell a visible grid border (Table Grid equivalent via XML)."""
    from docx.oxml.ns import qn
    tblPr = tbl._tbl.tblPr
    # remove existing borders element if present
    for el in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(el)
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = tblPr.makeelement(qn('w:' + edge), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '0', qn('w:color'): '666666'})
        borders.append(el)
    tblPr.append(borders)

def keep_rows_together(tbl):
    """Prevent table rows from splitting across pages (cantSplit on each row)."""
    from docx.oxml.ns import qn
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = trPr.makeelement(qn('w:cantSplit'), {})
        trPr.append(cant)

def render_practice(doc, s):
    if s.get("style") == "matching":
        doc.add_page_break()
    heading(doc, s.get("title", "Practice"), size=11.5, before=12)
    items = s.get("items", [])
    # matching: one table with two columns (col A / col B) + answer blanks
    if s.get("style") == "matching":
        col_a = [it.get("left", "") for it in items]
        col_b = [it.get("right", "") for it in items]
        n = len(items)
        tbl = doc.add_table(rows=n + 1, cols=3)
        tbl.autofit = False
        add_table_borders(tbl)
        keep_rows_together(tbl)
        for row in tbl.rows:
            for j, cell in enumerate(row.cells):
                cell.width = [Inches(0.5), Inches(2.9), Inches(3.6)][j]
        for j, h in enumerate(("", "Column A", "Column B")):
            cell = tbl.cell(0, j)
            shade_cell(cell, "EEF2FF")
            rr = cell.paragraphs[0].add_run(h)
            rr.font.bold = True; rr.font.size = Pt(9)
        for i in range(n):
            c0 = tbl.cell(i + 1, 0)
            c0.paragraphs[0].add_run("___").font.size = Pt(9.5)
            c1 = tbl.cell(i + 1, 1)
            add_chem_text(c1.paragraphs[0], col_a[i], size=9, color=SLATE)
            c2 = tbl.cell(i + 1, 2)
            add_chem_text(c2.paragraphs[0], col_b[i], size=9, color=SLATE)
        return
    for i, item in enumerate(items):
        q = item.get("q", "")
        t = item.get("type", "short")
        qnum = s.get("num_start", 1) + i
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        add_chem_text(p, f"{qnum}.  ", size=9.8, bold=True, color=INDIGO)
        add_chem_text(p, q, size=9.8)
        if t == "mc":
            for j, ch in enumerate(item.get("choices", [])):
                para(doc, f"   ({LETTERS[j]})  {ch}", size=9.5, color=SLATE, after=0)
            para(doc, "   Answer: ______", size=9, color=GRAY, after=5)
        elif t == "tf":
            para(doc, "   TRUE / FALSE", size=9.5, color=GRAY, italic=True, after=0)
            para(doc, "   If FALSE, correct the statement: _______________________________",
                 size=9, color=GRAY, after=5)
        elif t == "calc":
            # show-your-work box (2 blank lines)
            para(doc, "   Show your work:", size=9, color=GRAY, italic=True, after=1)
            for _ in range(2):
                para(doc, "", size=8, after=6)
        else:  # short — with writing lines
            for _ in range(2):
                para(doc, "   _________________________________________________________________",
                     size=9.5, color=SLATE, after=3)

def _render_vocabulary_with_definitions(doc, s):
    heading(doc, s.get("title", "Key Terms"), size=11.5, before=12)
    tbl = doc.add_table(rows=len(s.get("items", [])) + 1, cols=2)
    tbl.autofit = False
    widths = [Inches(2.2), Inches(4.8)]
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    for j, h in enumerate(("Term", "Definition")):
        cell = tbl.cell(0, j)
        shade_cell(cell, "EEF2FF")
        rr = cell.paragraphs[0].add_run(h)
        rr.font.bold = True; rr.font.size = Pt(8.5)
    for i, (term, definition) in enumerate(s.get("items", [])):
        c0 = tbl.cell(i + 1, 0)
        add_chem_text(c0.paragraphs[0], term, size=9, bold=True)
        c1 = tbl.cell(i + 1, 1)
        add_chem_text(c1.paragraphs[0], definition, size=9)


def render_figure(doc, s):
    """Embed an image with a caption (used for GHS grid, lab equipment, etc.)."""
    from docx.shared import Inches as _In
    img_rel = s.get("image", "")
    img_path = os.path.join(BASE, img_rel)
    if not os.path.exists(img_path):
        para(doc, f"[missing image: {img_rel}]", size=9, color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run()
    run.add_picture(img_path, width=_In(6.2))
    cap = s.get("caption", "")
    if cap:
        para(doc, cap, size=8.5, color=GRAY, italic=True, after=6,
             align=WD_ALIGN_PARAGRAPH.CENTER)


def render_review(doc, s):
    """Unit Review: end-of-unit multiple choice (Regents style)."""
    doc.add_page_break()
    heading(doc, s.get("title", "Unit Review"), size=13, before=0, color=INDIGO)
    para(doc, "Circle the number of the best answer for each question.",
         size=9.5, color=GRAY, italic=True, after=6)
    items = s.get("items", [])
    for i, item in enumerate(items):
        qnum = s.get("num_start", 1) + i
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        add_chem_text(p, f"{qnum}.  ", size=9.8, bold=True, color=INDIGO)
        add_chem_text(p, item.get("q", ""), size=9.8)
        for j, ch in enumerate(item.get("choices", [])):
            para(doc, f"   ({j+1})  {ch}", size=9.5, color=SLATE, after=0)
        para(doc, "", size=5, after=3)

def render_vocabulary(doc, s):
    """Render either a term/definition table or Regents-style writing prompts.

    Older product JSON stores vocabulary as ``[term, definition]`` pairs, while
    the Chapter 1 schema stores plain terms for students to define.  Supporting
    both shapes here keeps every committed product reproducible.
    """
    items = s.get("items", [])
    if items and all(isinstance(item, (list, tuple)) and len(item) >= 2
                     for item in items):
        return _render_vocabulary_with_definitions(doc, s)

    heading(doc, s.get("title", "Vocabulary"), size=11.5, before=12)
    for i, term in enumerate(items):
        if isinstance(term, list):
            term = term[0]
        para(doc, term, size=9.8, bold=True, color=INDIGO, after=1)
        for _ in range(1):
            para(doc, "   _________________________________________________________________",
                 size=9.5, color=SLATE, after=4)


# ---------- Regents-style renderers ----------

def render_assess_yourself(doc, s):
    """Assess Yourself checklist box (gray shaded, Vocab/Lesson 1-N lines)."""
    items = s.get("items", [])
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(0.4)
    tbl.columns[1].width = Inches(6.6)
    cell = tbl.cell(0, 0)
    shade_cell(cell, "D9D9D9")
    p0 = cell.paragraphs[0]
    r0 = p0.add_run("")
    r0.font.size = Pt(8)
    cell2 = tbl.cell(0, 1)
    shade_cell(cell2, "D9D9D9")
    p = cell2.paragraphs[0]
    r = p.add_run("Assess Yourself:")
    r.font.bold = True; r.font.size = Pt(10)
    for it in items:
        p2 = cell2.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(it + ": " + "____" * 12)
        r2.font.size = Pt(9.5)

def render_lesson(doc, s):
    """Regents-style lesson: ALL-CAPS header + objective box + concept + example + practice."""
    if s.get("page_break", True):
        doc.add_page_break()
    # ALL CAPS lesson header + divider line
    title = s.get("title", "Lesson").upper()
    heading(doc, title, size=14, before=0, color=INDIGO)
    # divider line
    pdiv = doc.add_paragraph()
    pdiv.paragraph_format.space_after = Pt(6)
    pPr = pdiv._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pPr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '12', qn('w:space'): '1', qn('w:color'): '444444'})
    pBdr.append(bottom)
    pPr.append(pBdr)
    # objective box (gray shaded)
    obj = s.get("objective", "")
    if obj:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(7.0)
        cell = tbl.cell(0, 0)
        shade_cell(cell, "EFEFEF")
        p0 = cell.paragraphs[0]
        r0 = p0.add_run("Objective:")
        r0.font.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = INDIGO
        if isinstance(obj, list):
            for b in obj:
                pb = cell.add_paragraph()
                pb.paragraph_format.space_after = Pt(1)
                rb = pb.add_run("•  " + b)
                rb.font.size = Pt(9.5)
        else:
            pb = cell.add_paragraph()
            rb = pb.add_run("•  " + obj)
            rb.font.size = Pt(9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    # concept paragraphs
    for para_text in s.get("concept", []):
        para(doc, para_text, size=9.5, after=4)
    # worked example (steps + OMML)
    ex = s.get("example")
    if ex:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Worked Example")
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = ORANGE
        ppr = doc.add_paragraph()
        ppr.paragraph_format.space_after = Pt(2)
        add_chem_text(ppr, ex.get("problem", ""), size=9.5)
        for step in ex.get("steps", []):
            if isinstance(step, str):
                step = {"text": step}
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.15)
            r = p.add_run(step.get("label", "Step") + ": ")
            r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INDIGO
            add_chem_text(p, step.get("text", ""), size=9.5)
            eq = step.get("eq")
            if eq:
                peq = doc.add_paragraph()
                peq.paragraph_format.left_indent = Inches(0.35)
                peq.paragraph_format.space_after = Pt(3)
                add_chem_text(peq, eq, size=10)
        if ex.get("answer"):
            pa = doc.add_paragraph()
            pa.paragraph_format.space_before = Pt(4)
            pa.paragraph_format.space_after = Pt(8)
            r = pa.add_run("Answer: ")
            r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = ORANGE
            add_chem_text(pa, ex.get("answer", ""), size=9.5)
    # embedded practice
    items = s.get("practice", [])
    if items:
        para(doc, "Check Your Understanding", size=10.5, bold=True, color=INDIGO, before=6, after=4)
        for i, item in enumerate(items):
            qnum = i + 1
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
            add_chem_text(p, f"{qnum}.  ", size=9.8, bold=True, color=INDIGO)
            add_chem_text(p, item.get("q", ""), size=9.8)
            img = item.get("image")
            if img:
                ip = doc.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                from docx.shared import Inches as _In
                ip.add_run().add_picture(os.path.join(BASE, img), width=_In(2.5))
            if item.get("type") == "mc":
                letters = item.get("option_style", "A-D")
                if letters == "1-4":
                    for j, ch in enumerate(item.get("choices", [])):
                        para(doc, f"   ({j+1})  {ch}", size=9.5, color=SLATE, after=0)
                else:
                    for j, ch in enumerate(item.get("choices", [])):
                        para(doc, f"   ({LETTERS[j]})  {ch}", size=9.5, color=SLATE, after=0)
                para(doc, "   Answer: ______", size=9, color=GRAY, after=5)
            else:
                for _ in range(2):
                    para(doc, "   _________________________________________________________________",
                         size=9.5, color=SLATE, after=3)

def render_fill_blank(doc, s):
    """Fill-in-the-blank sentences (_____ in running text)."""
    heading(doc, s.get("title", "Complete the Statements"), size=11.5, before=12)
    for i, sentence in enumerate(s.get("items", [])):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_chem_text(p, f"{i+1}.  ", size=9.8, bold=True, color=INDIGO)
        add_chem_text(p, sentence, size=9.8)

def render_data_table(doc, s):
    """Data table: header gray, some cells blank for students to fill."""
    doc.add_page_break()  # keep header + rows on one page
    heading(doc, s.get("title", "Complete the Table"), size=11.5, before=0)
    headers = s.get("headers", [])
    rows = s.get("rows", [])
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    add_table_borders(tbl)
    keep_rows_together(tbl)
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        shade_cell(cell, "D9D9D9")
        rr = cell.paragraphs[0].add_run(h)
        rr.font.bold = True; rr.font.size = Pt(8.5)
        cell.width = Inches(max(0.35, 6.8 / len(headers)))
    ncols = len(headers)
    for i, row in enumerate(rows):
        # pad short rows so every cell exists
        padded = list(row) + [""] * (ncols - len(row))
        for j, val in enumerate(padded[:ncols]):
            cell = tbl.cell(i + 1, j)
            cell.width = Inches(max(0.35, 6.8 / ncols))
            rr = cell.paragraphs[0].add_run("" if val == "" else str(val))
            rr.font.size = Pt(8.5)

def render_regents_practice(doc, s):
    """Regents Practice: exam-style MCQ (options 1-4) + open response."""
    doc.add_page_break()
    heading(doc, s.get("title", "Regents Practice"), size=13, before=0, color=INDIGO)
    para(doc, s.get("note", "Answer the following questions in the style of the NY State Regents exam."),
         size=9, color=GRAY, italic=True, after=6)
    for i, item in enumerate(s.get("items", [])):
        qnum = s.get("num_start", 1) + i
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        add_chem_text(p, f"{qnum}.)  ", size=9.8, bold=True, color=INDIGO)
        add_chem_text(p, item.get("q", ""), size=9.8)
        img = item.get("image")
        if img:
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.shared import Inches as _In
            ip.add_run().add_picture(os.path.join(BASE, img), width=_In(2.5))
        choices = item.get("choices", [])
        if choices:
            for j, ch in enumerate(choices):
                para(doc, f"   ({j+1})  {ch}", size=9.5, color=SLATE, after=0)
        else:
            # open response — lines
            for _ in range(2):
                para(doc, "   _________________________________________________________________",
                     size=9.5, color=SLATE, after=3)
        para(doc, "", size=4, after=2)

def render_answer_key(doc, s):
    doc.add_page_break()
    heading(doc, s.get("title", "Answer Key"), size=13, before=0)
    para(doc, s.get("note", ""), size=9.5, color=GRAY, italic=True, after=8)
    # collect answers from practice sections
    for sec in s.get("practice_sections", []):
        heading(doc, sec.get("title", ""), size=10.5, before=8)
        for i, item in enumerate(sec.get("items", [])):
            ans = item.get("answer", "")
            if not ans:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{sec.get('num_start', 1) + i}.  ")
            r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INDIGO
            add_chem_text(p, ans, size=9.5, color=SLATE)
            if item.get("explain"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(3)
                p2.paragraph_format.left_indent = Inches(0.3)
                r2 = p2.add_run("Why: ")
                r2.font.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GRAY
                add_chem_text(p2, item["explain"], size=9, color=GRAY)

LETTERS = list("ABCD")

# ---------------------------------------------------------------- main
def build(data):
    doc = new_doc()
    add_page_border(doc)
    add_header_footer(doc, "CurioNest — " + data.get("title", ""))

    # pass 1: collect practice sections for answer key
    practice_sections = []
    for s in data.get("sections", []):
        if s.get("kind") == "practice":
            practice_sections.append(s)
    # assign question numbers sequentially
    num = 1
    for ps in practice_sections:
        ps["num_start"] = num
        num += len(ps.get("items", []))
    for s in data.get("sections", []):
        if s.get("kind") == "answer_key":
            s["practice_sections"] = practice_sections

    for s in data.get("sections", []):
        kind = s.get("kind")
        if kind == "header":
            render_header(doc, s)
        elif kind == "learning_objectives":
            render_learning_objectives(doc, s)
        elif kind == "concept":
            render_concept(doc, s)
        elif kind == "practice":
            render_practice(doc, s)
        elif kind == "figure":
            render_figure(doc, s)
        elif kind == "assess_yourself":
            render_assess_yourself(doc, s)
        elif kind == "lesson":
            render_lesson(doc, s)
        elif kind == "fill_blank":
            render_fill_blank(doc, s)
        elif kind == "data_table":
            render_data_table(doc, s)
        elif kind == "regents_practice":
            render_regents_practice(doc, s)
        elif kind == "review":
            render_review(doc, s)
        elif kind == "vocabulary":
            render_vocabulary(doc, s)
        elif kind == "answer_key":
            render_answer_key(doc, s)
        else:
            print(f"  (skip unknown kind: {kind})")
    return doc

def main():
    if len(sys.argv) < 2:
        sys.exit("usage: python build_mirror_pack.py mirror-json/b01-mirror.json [outstem]")
    src = sys.argv[1]
    data = json.load(open(src, encoding="utf-8"))
    stem = sys.argv[2] if len(sys.argv) > 2 else data.get("meta", {}).get("file_stem", "pack")
    docx_path = os.path.join(BASE, f"{stem}-docx.docx")
    doc = build(data)
    doc.save(docx_path)
    print(f"docx saved: {docx_path}")
    # Word COM → PDF
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_docx = os.path.abspath(docx_path)
        pdf_path = os.path.join(BASE, f"{stem}-docx.pdf")
        doc_obj = word.Documents.Open(abs_docx)
        doc_obj.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc_obj.Close(False)
        word.Quit()
        print(f"pdf saved: {pdf_path}")
    except Exception as e:
        print(f"PDF conversion failed: {e}")

if __name__ == "__main__":
    main()
