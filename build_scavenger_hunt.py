#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""CurioNest Scavenger Hunt / Station builder — stations with clues → answer sheet → docx → PDF.

Schema (mirror-json/hunt-*.json):
{
  "title": "...", "subtitle": "...",
  "meta": {"file_stem": "...", "duration": "30 minutes", "stations": 8},
  "sections": [
    {"kind":"header"},
    {"kind":"setup", "items": ["Print station cards...", "Hang cards around the room..."]},
    {"kind":"stations", "title":"Station Cards", "items":[
        {"num":1, "title":"Station 1 — Safety Symbols", "clue":"Identify the symbol shown...", "task":"Write the meaning...", "answer":"Corrosive — burns skin and eyes"},
        ...
    ]},
    {"kind":"answer_sheet", "title":"Student Answer Sheet", "note":"Print one per student"},
    {"kind":"answer_key"}   # auto-built from stations
  ]
}
"""
import json, os, re, sys

BASE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INDIGO = RGBColor(0x31, 0x2E, 0x81)
PURPLE = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x33, 0x41, 0x55)
GRAY = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0x9F, 0x12, 0x39)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x15, 0x80, 0x3D)
AMBER = RGBColor(0xB4, 0x53, 0x09)

def add_chem_text(par, text, size=10.5, bold=False, color=None, italic=False):
    from lxml import etree
    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    for seg in re.split(r"(\$[^$]+\$)", text):
        if not seg:
            continue
        if seg.startswith("$") and seg.endswith("$") and len(seg) > 2:
            try:
                from omml import latex_to_omml
                body = latex_to_omml(seg[1:-1])
                omath = etree.fromstring(
                    f'<m:oMath xmlns:m="{M_NS}">'
                    f'<m:r><m:rPr><m:sz m:val="{int(size * 2)}"/></m:rPr></m:r>{body}'
                    f"</m:oMath>")
                par._p.append(omath)
            except Exception:
                r = par.add_run(seg)
                r.font.size = Pt(size); r.font.bold = bold
                if color: r.font.color.rgb = color
            continue
        stack = []
        pos = 0
        for m in re.finditer(r"<(/?)(sub|sup|b|i|br)([^>]*)>", seg):
            if m.start() > pos:
                run = par.add_run(seg[pos:m.start()])
                run.font.size = Pt(size)
                run.font.bold = bold or "b" in stack
                run.font.italic = italic or "i" in stack
                if color: run.font.color.rgb = color
                if "sub" in stack: run.font.subscript = True
                if "sup" in stack: run.font.superscript = True
            closing, tag, _ = m.groups()
            if tag == "br":
                par.add_run().add_break()
            elif closing:
                if tag in stack: stack.remove(tag)
            else:
                stack.append(tag)
            pos = m.end()
        if pos < len(seg):
            run = par.add_run(seg[pos:])
            run.font.size = Pt(size)
            run.font.bold = bold or "b" in stack
            run.font.italic = italic or "i" in stack
            if color: run.font.color.rgb = color
            if "sub" in stack: run.font.subscript = True
            if "sup" in stack: run.font.superscript = True

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcPr.append(shd)

def cell_borders(cell, color="999999", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin, sec.bottom_margin = Inches(0.55), Inches(0.6)
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(2)
    return doc

def add_page_border(doc, color="4F46E5", sz="10"):
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), color)
        pgBorders.append(el)
    sectPr.append(pgBorders)

def add_header_footer(doc, brand_text):
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(brand_text)
    r.font.size = Pt(7.5); r.font.color.rgb = GRAY; r.font.name = "Georgia"
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = fp.add_run("© 2026 CurioNest · For classroom use only · Page ")
    r2.font.size = Pt(7.5); r2.font.color.rgb = GRAY
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    fp.add_run()._r.append(fld1)
    fp.add_run()._r.append(instr)
    fp.add_run()._r.append(fld2)

def para(doc, text, size=10, color=SLATE, bold=False, italic=False, after=3,
         align=WD_ALIGN_PARAGRAPH.LEFT, indent=None, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if before: p.paragraph_format.space_before = Pt(before)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    add_chem_text(p, text, size=size, bold=bold, color=color, italic=italic)
    return p

def heading(doc, text, size=12, color=INDIGO, bold=True, after=4, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    add_chem_text(p, text, size=size, bold=bold, color=color)
    return p

# ---------------------------------------------------------------- renderers
def render_header(doc, s):
    para(doc, "", size=4, after=0)
    para(doc, "CURIONEST", size=13, color=INDIGO, bold=True, after=4,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, s.get("title", ""), size=18, color=INDIGO, bold=True, after=3,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, s.get("subtitle", ""), size=10, color=GRAY, italic=True, after=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    meta = s.get("meta", {})
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [("Duration", meta.get("duration", "30 minutes")),
            ("Stations", str(meta.get("stations", "—"))),
            ("Students", meta.get("students", "Pairs or small groups")),
            ("Includes", "Station cards · Student answer sheet · Answer key")]
    for i, (k, v) in enumerate(rows):
        c0, c1 = tbl.cell(i, 0), tbl.cell(i, 1)
        shade_cell(c0, "EEF2FF")
        add_chem_text(c0.paragraphs[0], k, size=9, bold=True)
        add_chem_text(c1.paragraphs[0], v, size=9, color=SLATE)
    para(doc, "Made by a chemistry teacher — gets students moving and thinking.",
         size=8.5, color=GRAY, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

def render_setup(doc, s):
    heading(doc, "Teacher Setup", size=11.5, before=0)
    for it in s.get("items", []):
        para(doc, "•  " + it, size=9.8, indent=0.2, after=2)

def render_stations(doc, s):
    heading(doc, s.get("title", "Station Cards"), size=11.5, before=8)
    # 2 stations per row (landscape card table: 2 cols x ceil(n/2) rows)
    items = s.get("items", [])
    rows = (len(items) + 1) // 2
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Inches(3.6)
            cell_borders(cell, color="C7D2FE", sz="6")
    for idx, st in enumerate(items):
        cell = tbl.cell(idx // 2, idx % 2)
        shade_cell(cell, "F5F7FF")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"STATION {st.get('num', idx + 1)} — {st.get('title', '')}")
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INDIGO
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_chem_text(p2, st.get("clue", ""), size=9, color=SLATE)
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run("TASK: ")
        r3.font.bold = True; r3.font.size = Pt(9); r3.font.color.rgb = AMBER
        add_chem_text(p3, st.get("task", ""), size=9, color=SLATE)

def render_answer_sheet(doc, s):
    doc.add_page_break()
    heading(doc, s.get("title", "Student Answer Sheet"), size=12, before=0)
    para(doc, s.get("note", "Print one per student. Walk around and record your answer at each station."),
         size=9.5, color=GRAY, italic=True, after=8)
    n = s.get("n_stations", 8)
    tbl = doc.add_table(rows=n + 1, cols=3)
    tbl.autofit = False
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = Inches([0.8, 4.3, 2.0][j])
    for j, h in enumerate(("Station", "Answer", "Self-check ✓")):
        cell = tbl.cell(0, j)
        shade_cell(cell, "EEF2FF")
        rr = cell.paragraphs[0].add_run(h)
        rr.font.bold = True; rr.font.size = Pt(9)
    for i in range(n):
        c0 = tbl.cell(i + 1, 0)
        add_chem_text(c0.paragraphs[0], str(i + 1), size=9.5, bold=True)
        for j in (1, 2):
            add_chem_text(tbl.cell(i + 1, j).paragraphs[0], "", size=9)

def render_answer_key(doc, s, stations):
    doc.add_page_break()
    heading(doc, "Answer Key", size=13, before=0)
    para(doc, "For the teacher — station answers with explanations.",
         size=9.5, color=GRAY, italic=True, after=8)
    for st in stations:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"Station {st.get('num', '')} — {st.get('title', '')}: ")
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INDIGO
        add_chem_text(p, st.get("answer", ""), size=9.5, color=GREEN, bold=True)
        if st.get("explain"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            p2.paragraph_format.left_indent = Inches(0.3)
            r2 = p2.add_run("Why: ")
            r2.font.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GRAY
            add_chem_text(p2, st["explain"], size=9, color=GRAY)

# ---------------------------------------------------------------- main
def build(data):
    doc = new_doc()
    add_page_border(doc)
    add_header_footer(doc, "CurioNest — " + data.get("title", ""))
    stations = []
    for s in data.get("sections", []):
        if s.get("kind") == "stations":
            stations = s.get("items", [])
        if s.get("kind") == "header":
            s["meta"] = data.get("meta", {})
            s["meta"]["stations"] = len(stations)
        if s.get("kind") == "answer_sheet":
            s["n_stations"] = len(stations)
    for s in data.get("sections", []):
        kind = s.get("kind")
        if kind == "header":
            render_header(doc, s)
        elif kind == "setup":
            render_setup(doc, s)
        elif kind == "stations":
            render_stations(doc, s)
        elif kind == "answer_sheet":
            render_answer_sheet(doc, s)
        elif kind == "answer_key":
            render_answer_key(doc, s, stations)
        else:
            print(f"  (skip unknown kind: {kind})")
    return doc

def main():
    if len(sys.argv) < 2:
        sys.exit("usage: python build_scavenger_hunt.py mirror-json/hunt-x.json [outstem]")
    src = sys.argv[1]
    data = json.load(open(src, encoding="utf-8"))
    stem = sys.argv[2] if len(sys.argv) > 2 else data.get("meta", {}).get("file_stem", "hunt")
    docx_path = os.path.join(BASE, f"{stem}-docx.docx")
    build(data).save(docx_path)
    print(f"docx saved: {docx_path}")
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(docx_path))
        pdf_path = os.path.join(BASE, f"{stem}-docx.pdf")
        d.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        d.Close(False)
        word.Quit()
        print(f"pdf saved: {pdf_path}")
    except Exception as e:
        print(f"PDF conversion failed: {e}")

if __name__ == "__main__":
    main()
