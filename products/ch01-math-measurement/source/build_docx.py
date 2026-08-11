"""Build a stable, editable US Letter DOCX companion for Chapter 1."""

from pathlib import Path
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
PRODUCT = HERE.parent
ASSETS = PRODUCT / "assets"
OUT = PRODUCT / "output" / "buyer-files"
DATA = json.loads((HERE / "source.json").read_text(encoding="utf-8"))
COPYRIGHT = DATA["copyright"]
FONT = "Aptos"
NAVY, ORANGE, WASH, NOTE, MUTED, WHITE = "203B4D", "C86543", "F8F5EE", "F3E5BD", "6D7478", "FFFFFF"


def run_style(run, size=8, bold=False, color="202B33", italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT); rpr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); ppr.append(shd)


def add_p(doc, text="", size=8, bold=False, color="202B33", after=2, align=None, fill=None, keep=False):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.0
    if align is not None: p.alignment = align
    if fill: shade_paragraph(p, fill); p.paragraph_format.left_indent = Inches(.07); p.paragraph_format.right_indent = Inches(.07)
    if keep: p.paragraph_format.keep_with_next = True
    run_style(p.add_run(str(text)), size, bold, color)
    return p


def add_page_header(doc, page, key):
    p = add_p(doc, f"{page['number']}   {page['title']}", 16, True, NAVY, 0, keep=True)
    p.paragraph_format.top_border = None
    add_p(doc, page["subtitle"], 7.5, False, MUTED, 2, keep=True)
    if key:
        add_p(doc, "TEACHER KEY", 7, True, ORANGE, 3, WD_ALIGN_PARAGRAPH.RIGHT, NOTE)
    else:
        add_p(doc, "Name: ____________________    Date: __________    Class Period: ______", 6.5, False, MUTED, 3, WD_ALIGN_PARAGRAPH.RIGHT)


def add_special(doc, page):
    kinds = {b["type"]: b for b in page["blocks"]}
    if "targets" in kinds:
        add_p(doc, "LEARNING TARGET  Record complete measurements, select a tool from evidence, and explain why units and precision matter.", 7.2, True, NAVY, 3, fill=WASH)
    if "image_row" in kinds:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0); shade_paragraph(p, WASH)
        for item in kinds["image_row"]["assets"]:
            p.add_run("      ").add_picture(str(ASSETS / item["file"]), height=Inches(.58))
        add_p(doc, "Electronic balance                         Graduated cylinder                         Beaker", 6.8, True, NAVY, 2, WD_ALIGN_PARAGRAPH.CENTER, WASH)
    if "prefix_table" in kinds:
        add_p(doc, "kilo: k, 10^3     |     centi: c, 10^-2     |     milli: m, 10^-3     |     micro: µ, 10^-6     |     nano: n, 10^-9", 7.1, True, WHITE, 3, WD_ALIGN_PARAGRAPH.CENTER, NAVY)
    if "rule_box" in kinds:
        b = kinds["rule_box"]; add_p(doc, f"{b['title'].upper()}  {b['text']}", 7.2, True, NAVY, 3, fill=NOTE)
    if "photo_strip" in kinds:
        b = kinds["photo_strip"]
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
        p.add_run().add_picture(str(ASSETS / b["file"]), height=Inches(.72))
        add_p(doc, "AUTHENTIC MEASUREMENT CONTEXT  Read a concave meniscus at eye level. A narrow graduated vessel supports more precise volume readings than a beaker.", 7, True, NAVY, 3, fill=WASH)
    if "formula_box" in kinds:
        add_p(doc, "FORMULA TOOLBOX  density = mass / volume    |    percent error = |experimental - accepted| / accepted × 100", 7.2, True, NAVY, 3, fill=NOTE)
    if "case_box" in kinds:
        add_p(doc, "CASE FILE  " + kinds["case_box"]["text"], 7.1, True, NAVY, 3, fill=NOTE)


def add_questions(doc, page, key):
    items = next(b["items"] for b in page["blocks"] if b["type"] == "questions")
    size = 6.8 if len(items) > 5 else 7.4
    for item in items:
        add_p(doc, f"{item['id']}. {item['prompt']}", size, True, "202B33", 0, fill=WASH, keep=True)
        if key:
            add_p(doc, "ANSWER: " + item["answer"], size-.2, True, NAVY, 3, fill=NOTE)
        else:
            add_p(doc, "________________________________________________________________________________________\n________________________________________________________________________________________", 5.4, False, "B8B6B0", 3)


def make_docx():
    doc = Document(); sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(.45); sec.bottom_margin = Inches(.48); sec.left_margin = Inches(.58); sec.right_margin = Inches(.58)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_style(footer.add_run(f"{COPYRIGHT}  |  EDITABLE STUDENT + KEY"), 6.3, False, MUTED)
    normal = doc.styles["Normal"]; normal.font.name = FONT; normal.font.size = Pt(8)
    props = doc.core_properties; props.author = "CurioNest"; props.creator = "CurioNest"; props.title = DATA["title"] + " - Editable"; props.subject = COPYRIGHT
    first = True
    for key in (False, True):
        for page in DATA["pages"]:
            if not first: doc.add_page_break()
            first = False; add_page_header(doc, page, key); add_special(doc, page); add_questions(doc, page, key)
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / "CurioNest_CH01_Math_and_Measurement_Editable.docx"; doc.save(path); print(path)


if __name__ == "__main__": make_docx()
