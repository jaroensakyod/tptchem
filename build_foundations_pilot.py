#!/usr/bin/env python3
"""Build the Chemistry Foundations pilot and its PDF-ready assets.

Scientific diagrams remain deterministic, while standardized equipment and
GHS visuals come from the project's verified Bioicons source set.
"""

from __future__ import annotations

import math
import os
import shutil
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "products" / "chemistry-foundations-pilot"
ASSET_DIR = OUT / "assets"
DOCX_PATH = OUT / "product-editable.docx"
BIOICON_EQUIPMENT_DIR = ROOT / "figures" / "bioicons" / "verified-equipment"
BIOICON_GHS_DIR = ROOT / "figures" / "bioicons" / "verified-ghs"

NAVY = "0B1F3A"
TEAL = "148C7E"
TEAL_DARK = "0F6D63"
AMBER = "F4A261"
RED = "C44536"
BLUE = "337AB7"
INK = "243447"
GRAY = "607080"
LIGHT = "EAF7F4"
PALE = "F5F8FA"
WHITE = "FFFFFF"
LINE = "B8C6D1"

PAGE_W, PAGE_H = 1275, 1650
ART_W = 1600


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded(draw, box, radius=24, fill=WHITE, outline=LINE, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


# These three schematics support measurement reasoning. Standardized equipment
# identification elsewhere in the product uses the verified Bioicons asset set.
def draw_beaker(d, b, color=BLUE):
    x1,y1,x2,y2=b; w=x2-x1; h=y2-y1
    pts=[(x1+.25*w,y1+.18*h),(x1+.75*w,y1+.18*h),(x1+.68*w,y1+.82*h),(x1+.32*w,y1+.82*h)]
    d.polygon(pts, fill="#DDEFF8", outline=f"#{color}", width=5)
    d.line((x1+.30*w,y1+.58*h,x1+.70*w,y1+.58*h), fill="#5CB8D8", width=5)
    for frac in (.35,.45,.55): d.line((x1+.34*w,y1+frac*h,x1+.44*w,y1+frac*h),fill=f"#{color}",width=3)


def draw_cylinder(d,b,color=BLUE):
    x1,y1,x2,y2=b; w=x2-x1; h=y2-y1
    d.rounded_rectangle((x1+.39*w,y1+.12*h,x1+.61*w,y1+.82*h),radius=8,fill="#E5F3FA",outline=f"#{color}",width=5)
    d.line((x1+.27*w,y1+.84*h,x1+.73*w,y1+.84*h),fill=f"#{color}",width=6)
    for i in range(7):
        y=y1+(.22+i*.075)*h; d.line((x1+.40*w,y,x1+(.50 if i%2 else .55)*w,y),fill=f"#{color}",width=3)
    d.arc((x1+.40*w,y1+.53*h,x1+.60*w,y1+.60*h),0,180,fill="#3FA7CF",width=4)


def draw_balance(d,b,color=NAVY):
    x1,y1,x2,y2=b; w=x2-x1; h=y2-y1
    d.rounded_rectangle((x1+.18*w,y1+.38*h,x1+.82*w,y1+.78*h),radius=16,fill="#E8EDF2",outline=f"#{color}",width=5)
    d.ellipse((x1+.27*w,y1+.20*h,x1+.73*w,y1+.45*h),fill="#C9D4DD",outline=f"#{color}",width=4)
    d.rounded_rectangle((x1+.35*w,y1+.55*h,x1+.65*w,y1+.69*h),radius=6,fill="#C9F2E7",outline=f"#{TEAL}",width=3)
    d.text((x1+.50*w,y1+.62*h),"0.00 g",font=font(max(12,int(.08*w)),True),fill=f"#{NAVY}",anchor="mm")


EQUIPMENT = [
    ("Beaker", "Holds/mixes liquids; approximate volume"),
    ("Erlenmeyer flask", "Swirls mixtures with less splashing"),
    ("Graduated cylinder", "Measures liquid volume accurately"),
    ("Test tube", "Runs small-scale reactions"),
    ("Buret", "Delivers precise variable volumes"),
    ("Dropper pipet", "Transfers small amounts dropwise"),
    ("Funnel", "Guides pouring or supports filtration"),
    ("Electronic balance", "Measures mass; tare before use"),
    ("Hot plate", "Heats without an open flame"),
    ("Wash bottle", "Dispenses distilled/deionized water"),
    ("Ring stand", "Supports clamps and apparatus"),
    ("Crucible tongs", "Handles hot crucibles"),
]

EQUIPMENT_ICONS = {
    "Beaker": "beaker.png",
    "Erlenmeyer flask": "erlenmeyer-flask.png",
    "Graduated cylinder": "graduated-cylinder.png",
    "Test tube": "test-tube.png",
    "Buret": "buret.png",
    "Dropper pipet": "pipette.png",
    "Funnel": "funnel.png",
    "Electronic balance": "electronic-balance.png",
    "Hot plate": "hot-plate.png",
    "Wash bottle": "wash-bottle.png",
    "Ring stand": "ring-stand-servier.png",
    "Crucible tongs": "lab-tongs.png",
}


def paste_asset_icon(canvas: Image.Image, icon_path: Path, box) -> None:
    """Paste a sourced transparent icon centered inside a target box."""
    if not icon_path.exists():
        raise FileNotFoundError(f"Missing verified icon: {icon_path}")
    x1, y1, x2, y2 = (int(value) for value in box)
    icon = Image.open(icon_path).convert("RGBA")
    alpha_box = icon.getbbox()
    if alpha_box:
        icon = icon.crop(alpha_box)
    icon.thumbnail((x2-x1, y2-y1), Image.Resampling.LANCZOS)
    x = x1 + (x2-x1-icon.width)//2
    y = y1 + (y2-y1-icon.height)//2
    canvas.paste(icon,(x,y),icon)


def equipment_grid(path: Path, items, title: str, numbered=False, show_functions=True):
    im = Image.new("RGB", (1600, 1050), "white"); d=ImageDraw.Draw(im)
    d.text((800,52),title,font=font(42,True),fill=f"#{NAVY}",anchor="mm")
    cols = 3 if len(items) <= 6 else 4
    rows = math.ceil(len(items) / cols)
    gap=22; left=42; top=95
    cw=(1600-left*2-gap*(cols-1))//cols
    ch=(920-gap*(rows-1))//rows
    for i,(name,func) in enumerate(items):
        r=i//cols; c=i%cols; x=left+c*(cw+gap); y=top+r*(ch+gap)
        rounded(d,(x,y,x+cw,y+ch),22,fill="FFFFFF",outline="C8D5DE",width=3)
        icon_h = int(ch * (0.53 if show_functions else 0.64))
        paste_asset_icon(im,BIOICON_EQUIPMENT_DIR/EQUIPMENT_ICONS[name],(x+int(cw*.14),y+18,x+int(cw*.86),y+18+icon_h))
        label=f"{i+1}." if numbered else name
        label_y = y + int(ch * (0.68 if show_functions else 0.77))
        d.text((x+cw/2,label_y),label,font=font(24,True),fill=f"#{NAVY}",anchor="mm")
        if show_functions:
            words=func.split(); lines=[]; current=""
            for word in words:
                if len(current)+len(word)+1>34: lines.append(current); current=word
                else: current=(current+" "+word).strip()
            if current: lines.append(current)
            d.multiline_text((x+cw/2,y+int(ch*.82)),"\n".join(lines),font=font(17),fill=f"#{GRAY}",anchor="mm",align="center",spacing=3)
    im.save(path,quality=95)


def make_ghs_grid(path: Path):
    items=[
        ("Exploding bomb","Explosives; some self-reactives/peroxides","ghs-explosive.png"),("Flame","Flammable and related fire hazards","ghs-flammable.png"),("Oxidizer","Oxidizing gases, liquids, or solids","ghs-oxidizing.png"),
        ("Gas cylinder","Gases under pressure","ghs-gas-under-pressure.png"),("Corrosion","Skin burns/eye damage; corrosive to metals","ghs-corrosive.png"),("Skull","Fatal or toxic acute toxicity","ghs-acute-toxicity.png"),
        ("Exclamation","Irritant or harmful effects","ghs-health-hazard.png"),("Health hazard","Cancer, organ, respiratory, aspiration hazards","ghs-serious-health-hazard.png"),("Environment","Aquatic toxicity; not OSHA-mandated","ghs-hazardous-environment.png"),
    ]
    im=Image.new("RGB",(1600,1080),"white"); d=ImageDraw.Draw(im)
    d.text((800,48),"GHS Hazard Pictograms",font=font(43,True),fill=f"#{NAVY}",anchor="mm")
    for i,(name,desc,filename) in enumerate(items):
        r=i//3;c=i%3;x=45+c*515;y=95+r*315
        rounded(d,(x,y,x+485,y+290),24,fill="FFFFFF",outline="D7E0E6",width=3)
        icon_path = BIOICON_GHS_DIR / filename
        if not icon_path.exists():
            raise FileNotFoundError(f"Missing verified GHS icon: {icon_path}")
        icon = Image.open(icon_path).convert("RGBA")
        icon.thumbnail((210,210),Image.Resampling.LANCZOS)
        im.paste(icon,(x+15+(210-icon.width)//2,y+15+(210-icon.height)//2),icon)
        d.text((x+245,y+70),name,font=font(24,True),fill=f"#{NAVY}")
        words=desc.split();lines=[];cur=""
        for word in words:
            if len(cur)+len(word)+1>25:lines.append(cur);cur=word
            else:cur=(cur+" "+word).strip()
        if cur:lines.append(cur)
        d.multiline_text((x+245,y+112),"\n".join(lines),font=font(18),fill=f"#{GRAY}",spacing=5)
    im.save(path)


def make_safety_scenarios(path: Path):
    """Create a text-based hazard sort instead of decorative character art."""
    scenarios = [
        ("1", "EYE PROTECTION", "A student begins work with splash goggles resting on the forehead."),
        ("2", "FOOD & DRINK", "A drink cup sits beside the lab notebook and glassware."),
        ("3", "LABELING", "A beaker contains blue liquid but has no chemical label."),
        ("4", "HEAT", "Loose paper is placed beside a lit laboratory burner."),
        ("5", "BROKEN GLASS", "Broken glass is left on the bench after cleanup begins."),
        ("6", "HOUSEKEEPING", "A backpack blocks the aisle beside the lab station."),
        ("7", "CHEMICALS", "An open reagent bottle is left unattended on the bench."),
        ("8", "PERSONAL SAFETY", "Long hair hangs loose near the work area and heat source."),
    ]
    im=Image.new("RGB",(1600,900),"#F6FAFC");d=ImageDraw.Draw(im)
    d.rectangle((0,0,1600,130),fill=f"#{NAVY}")
    d.text((70,48),"LAB SAFETY DETECTIVE",font=font(46,True),fill="white",anchor="lm")
    d.text((70,98),"Eight classroom snapshots • Identify the unsafe choice before writing the safer action.",font=font(22),fill="#DCE8F2",anchor="lm")
    cols, rows, gap = 4, 2, 22
    left, top = 45, 165
    cw = (1600-left*2-gap*(cols-1))//cols
    ch = 315
    for i,(num,tag,text) in enumerate(scenarios):
        r,c=divmod(i,cols);x=left+c*(cw+gap);y=top+r*(ch+gap)
        rounded(d,(x,y,x+cw,y+ch),24,fill=WHITE,outline="C8D5DE",width=3)
        d.ellipse((x+22,y+22,x+82,y+82),fill=f"#{TEAL}")
        d.text((x+52,y+52),num,font=font(28,True),fill="white",anchor="mm")
        d.text((x+102,y+52),tag,font=font(18,True),fill=f"#{TEAL_DARK}",anchor="lm")
        words=text.split();lines=[];cur=""
        for word in words:
            if len(cur)+len(word)+1>31:lines.append(cur);cur=word
            else:cur=(cur+" "+word).strip()
        if cur:lines.append(cur)
        d.multiline_text((x+28,y+115),"\n".join(lines),font=font(22),fill=f"#{INK}",spacing=10)
        d.line((x+28,y+268,x+cw-28,y+268),fill=f"#{LINE}",width=3)
        d.text((x+28,y+288),"What should happen instead?",font=font(16,True),fill=f"#{GRAY}")
    im.save(path)


def make_label(path: Path):
    im=Image.new("RGB",(1600,820),"white");d=ImageDraw.Draw(im)
    rounded(d,(70,70,1530,750),28,fill="FFFFFF",outline=NAVY,width=5)
    d.rectangle((70,70,1530,170),fill=f"#{NAVY}")
    d.text((120,120),"SAMPLE CHEMICAL LABEL",font=font(40,True),fill="white",anchor="lm")
    d.text((120,230),"ACETONE",font=font(48,True),fill=f"#{NAVY}")
    d.text((120,290),"Signal word: DANGER",font=font(30,True),fill=f"#{RED}")
    paste_asset_icon(im,BIOICON_GHS_DIR/"ghs-flammable.png",(115,335,405,625))
    paste_asset_icon(im,BIOICON_GHS_DIR/"ghs-health-hazard.png",(420,335,710,625))
    bullets=["Highly flammable liquid and vapor.","Causes serious eye irritation.","Keep away from heat, sparks, and open flame.","Wear eye protection; use with adequate ventilation."]
    y=220
    for i,t in enumerate(bullets):
        d.ellipse((780,y+8,798,y+26),fill=f"#{TEAL}");d.text((820,y),t,font=font(24),fill=f"#{INK}");y+=72
    d.text((780,550),"Supplier: Example Chemical Co. | Emergency: 555-0100",font=font(21),fill=f"#{GRAY}")
    d.text((780,610),"Read the SDS before use. This fictional label is for instruction only.",font=font(20,True),fill=f"#{NAVY}")
    im.save(path)


def make_measurement(path: Path):
    im=Image.new("RGB",(1600,880),"white");d=ImageDraw.Draw(im)
    d.text((800,48),"Choose the Tool - Read the Scale - Report the Unit",font=font(42,True),fill=f"#{NAVY}",anchor="mm")
    panels=[("BEAKER","Approximate volume",draw_beaker),("GRADUATED CYLINDER","Measured volume",draw_cylinder),("ELECTRONIC BALANCE","Mass after taring",draw_balance)]
    for i,(title,sub,fn) in enumerate(panels):
        x=55+i*515;rounded(d,(x,100,x+480,815),24,fill="FFFFFF",outline="C8D5DE",width=3)
        d.text((x+240,145),title,font=font(28,True),fill=f"#{NAVY}",anchor="mm")
        d.text((x+240,188),sub,font=font(21),fill=f"#{GRAY}",anchor="mm")
        fn(d,(x+80,220,x+400,560))
        if i==1:
            d.line((x+370,310,x+445,310),fill=f"#{RED}",width=5);d.text((x+445,290),"eye level",font=font(18,True),fill=f"#{RED}",anchor="ra")
            d.text((x+240,610),"Read bottom of meniscus",font=font(22,True),fill=f"#{TEAL_DARK}",anchor="mm")
        elif i==2:
            d.text((x+240,610),"Container on -> TARE -> sample",font=font(21,True),fill=f"#{TEAL_DARK}",anchor="mm")
        else:
            d.text((x+240,610),"Do not claim cylinder-level precision",font=font(19,True),fill=f"#{TEAL_DARK}",anchor="mm")
        d.text((x+240,720),"Value: __________    Unit: ______",font=font(22),fill=f"#{INK}",anchor="mm")
    im.save(path)


def make_meniscus_practice(path: Path):
    im=Image.new("RGB",(1600,760),"white");d=ImageDraw.Draw(im)
    d.text((800,48),"Read Each Graduated Cylinder",font=font(42,True),fill=f"#{NAVY}",anchor="mm")
    configs=[("A",32.6,30,35,1),("B",18.4,15,20,1),("C",46.5,40,50,2)]
    for i,(label,val,lo,hi,minor) in enumerate(configs):
        x=90+i*510;rounded(d,(x,100,x+430,700),24,fill="FFFFFF",outline="C8D5DE",width=3)
        d.text((x+215,145),label,font=font(34,True),fill=f"#{NAVY}",anchor="mm")
        cx=x+215;top=190;bottom=545
        d.rounded_rectangle((cx-78,top,cx+78,bottom),radius=20,fill="#EFF8FC",outline=f"#{BLUE}",width=5)
        steps=int((hi-lo)/minor)
        for s in range(steps+1):
            y=bottom-45-s*(bottom-top-90)/steps
            d.line((cx-78,y,cx-(15 if s%5 else -15),y),fill=f"#{NAVY}",width=3)
            if s%5==0:d.text((cx-95,y),str(lo+s*minor),font=font(17),fill=f"#{GRAY}",anchor="rm")
        yval=bottom-45-(val-lo)/(hi-lo)*(bottom-top-90)
        d.rectangle((cx-74,yval,cx+74,bottom-10),fill="#8BD2E8")
        d.arc((cx-74,yval-16,cx+74,yval+18),0,180,fill=f"#{BLUE}",width=5)
        d.text((x+215,625),"Reading: ______ mL",font=font(24,True),fill=f"#{INK}",anchor="mm")
    im.save(path)


def make_cover(path: Path):
    im=Image.new("RGB",(PAGE_W,PAGE_H),f"#{PALE}");d=ImageDraw.Draw(im)
    d.rectangle((0,0,PAGE_W,340),fill=f"#{NAVY}")
    d.text((85,78),"CURIONEST CHEMISTRY",font=font(29,True),fill=f"#{AMBER}")
    d.text((85,160),"CHEMISTRY\nFOUNDATIONS",font=font(66,True),fill="white",spacing=4)
    d.text((88,303),"Lab Safety, Equipment & Measurement",font=font(30),fill="#DCE8F2")
    # visual cards
    cards=[
        ("GHS safety",BIOICON_GHS_DIR/"ghs-corrosive.png",TEAL),
        ("Graduated cylinder",BIOICON_EQUIPMENT_DIR/EQUIPMENT_ICONS["Graduated cylinder"],BLUE),
        ("Electronic balance",BIOICON_EQUIPMENT_DIR/EQUIPMENT_ICONS["Electronic balance"],AMBER),
        ("Erlenmeyer flask",BIOICON_EQUIPMENT_DIR/EQUIPMENT_ICONS["Erlenmeyer flask"],TEAL),
    ]
    for i,(name,icon_path,accent) in enumerate(cards):
        x=70+(i%2)*585;y=410+(i//2)*420
        rounded(d,(x,y,x+545,y+370),30,fill="FFFFFF",outline=accent,width=5)
        paste_asset_icon(im,icon_path,(x+120,y+35,x+425,y+265))
        d.text((x+272,y+320),name,font=font(26,True),fill=f"#{NAVY}",anchor="mm")
    rounded(d,(70,1285,1205,1480),26,fill=LIGHT,outline=TEAL,width=4)
    d.text((105,1330),"PRINT + EDITABLE",font=font(26,True),fill=f"#{TEAL_DARK}")
    d.text((105,1380),"Teacher guide | visual reference pages | stations | practical | quiz | full key",font=font(23),fill=f"#{INK}")
    d.text((105,1430),"Grades 9-11  •  3-5 class periods  •  US Letter",font=font(22,True),fill=f"#{NAVY}")
    im.save(path)


def make_listing_cover(path: Path):
    """Square TPT thumbnail that stays readable in search-result view."""
    im=Image.new("RGB",(1800,1800),f"#{PALE}");d=ImageDraw.Draw(im)
    d.rectangle((0,0,1800,560),fill=f"#{NAVY}")
    d.text((100,105),"CURIONEST CHEMISTRY",font=font(40,True),fill=f"#{AMBER}")
    d.text((100,225),"CHEMISTRY",font=font(88,True),fill="white")
    d.text((100,325),"FOUNDATIONS",font=font(88,True),fill="white")
    d.text((105,455),"LAB SAFETY • EQUIPMENT • MEASUREMENT",font=font(37,True),fill="#DCE8F2")
    cards=[
        ("Safety",BIOICON_GHS_DIR/"ghs-corrosive.png",TEAL),
        ("Equipment",BIOICON_EQUIPMENT_DIR/EQUIPMENT_ICONS["Erlenmeyer flask"],TEAL),
        ("Measurement",BIOICON_EQUIPMENT_DIR/EQUIPMENT_ICONS["Graduated cylinder"],BLUE),
    ]
    gap=35;left=70;cw=(1800-left*2-gap*2)//3
    for i,(label,icon_path,accent) in enumerate(cards):
        x=left+i*(cw+gap);rounded(d,(x,640,x+cw,1240),34,fill="FFFFFF",outline=accent,width=6)
        paste_asset_icon(im,icon_path,(x+100,700,x+cw-100,1080))
        d.text((x+cw/2,1165),label,font=font(38,True),fill=f"#{NAVY}",anchor="mm")
    rounded(d,(70,1320,1730,1715),34,fill=LIGHT,outline=TEAL,width=5)
    d.text((120,1385),"20-PAGE COMPLETE UNIT",font=font(46,True),fill=f"#{TEAL_DARK}")
    d.text((120,1470),"Teacher guide + student packet + visual references",font=font(31),fill=f"#{INK}")
    d.text((120,1530),"Stations + lab practical + 2-part quiz + full answer key",font=font(31),fill=f"#{INK}")
    d.text((120,1635),"PRINT  |  EDITABLE  |  GRADES 9-11",font=font(35,True),fill=f"#{NAVY}")
    im.save(path)


def prepare_assets():
    ASSET_DIR.mkdir(parents=True,exist_ok=True)
    make_cover(ASSET_DIR/"cover.png")
    make_listing_cover(OUT/"cover.png")
    make_safety_scenarios(ASSET_DIR/"safety-scene.png")
    make_ghs_grid(ASSET_DIR/"ghs-grid.png")
    make_label(ASSET_DIR/"sample-label.png")
    equipment_grid(ASSET_DIR/"equipment-reference-a.png",EQUIPMENT[:6],"Lab Equipment Reference - Set A")
    equipment_grid(ASSET_DIR/"equipment-reference-b.png",EQUIPMENT[6:12],"Lab Equipment Reference - Set B")
    equipment_grid(ASSET_DIR/"equipment-challenge.png",EQUIPMENT,"Equipment Identification Challenge",numbered=True,show_functions=False)
    make_measurement(ASSET_DIR/"measurement-guide.png")
    make_meniscus_practice(ASSET_DIR/"meniscus-practice.png")


def set_cell_fill(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr();shd=tcPr.find(qn("w:shd"))
    if shd is None:shd=OxmlElement("w:shd");tcPr.append(shd)
    shd.set(qn("w:fill"),fill)


def set_cell_margins(cell,top=90,start=120,bottom=90,end=120):
    tc=cell._tc;tcPr=tc.get_or_add_tcPr();tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:tcMar=OxmlElement("w:tcMar");tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn(f"w:{m}"))
        if node is None:node=OxmlElement(f"w:{m}");tcMar.append(node)
        node.set(qn("w:w"),str(v));node.set(qn("w:type"),"dxa")


def set_table_widths(table,widths_dxa):
    table.autofit=False;tblPr=table._tbl.tblPr
    tblW=tblPr.first_child_found_in("w:tblW")
    if tblW is None:tblW=OxmlElement("w:tblW");tblPr.append(tblW)
    tblW.set(qn("w:w"),str(sum(widths_dxa)));tblW.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for child in list(grid):grid.remove(child)
    for width in widths_dxa:
        col=OxmlElement("w:gridCol");col.set(qn("w:w"),str(width));grid.append(col)
    for row in table.rows:
        for cell,width in zip(row.cells,widths_dxa):
            tcW=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"),str(width));tcW.set(qn("w:type"),"dxa");set_cell_margins(cell)


def set_run(run,size=11,bold=False,color=INK,italic=False,font_name="Calibri"):
    run.font.name=font_name;run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),font_name);run._element.rPr.rFonts.set(qn("w:hAnsi"),font_name)
    run.font.size=Pt(size);run.bold=bold;run.italic=italic;run.font.color.rgb=RGBColor.from_string(color)


def style_doc(doc):
    sec=doc.sections[0];sec.page_width=Inches(8.5);sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=Inches(0.72);sec.left_margin=sec.right_margin=Inches(0.72)
    sec.header_distance=Inches(0.30);sec.footer_distance=Inches(0.32)
    normal=doc.styles["Normal"];normal.font.name="Calibri";normal.font.size=Pt(10.5);normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(6);normal.paragraph_format.line_spacing=1.12
    for sty,size,color,before,after in (("Title",28,NAVY,0,10),("Heading 1",16,NAVY,12,7),("Heading 2",13,TEAL_DARK,10,5),("Heading 3",11,INK,7,4)):
        s=doc.styles[sty];s.font.name="Calibri";s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(color)
        s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after);s.paragraph_format.keep_with_next=True
    # running header/footer
    hp=sec.header.paragraphs[0];hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("CURIONEST CHEMISTRY  |  FOUNDATIONS PILOT"),8,True,GRAY)
    fp=sec.footer.paragraphs[0];fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.add_run("Single-classroom license  |  © 2026 CurioNest  |  Page "),8,False,GRAY)
    fld=OxmlElement("w:fldSimple");fld.set(qn("w:instr"),"PAGE");fp._p.append(fld)


def add_title(doc,title,subtitle=None,kicker=None):
    if kicker:
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(3);set_run(p.add_run(kicker.upper()),9,True,TEAL_DARK)
    p=doc.add_paragraph(style="Title");p.paragraph_format.space_after=Pt(4);set_run(p.add_run(title),28,True,NAVY)
    if subtitle:
        p2=doc.add_paragraph();p2.paragraph_format.space_after=Pt(10);set_run(p2.add_run(subtitle),12,False,GRAY,True)


def add_para(doc,text,size=10.5,bold=False,color=INK,italic=False,after=6,align=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(after)
    if align is not None:p.alignment=align
    set_run(p.add_run(text),size,bold,color,italic);return p


def add_bullet(doc,text):
    p=doc.add_paragraph(style="List Bullet");p.paragraph_format.left_indent=Inches(.38);p.paragraph_format.first_line_indent=Inches(-.18);p.paragraph_format.space_after=Pt(3)
    set_run(p.add_run(text),10.2,False,INK);return p


def add_callout(doc,label,text,fill=LIGHT,accent=TEAL):
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;set_table_widths(t,[9360]);c=t.cell(0,0);set_cell_fill(c,fill)
    p=c.paragraphs[0];set_run(p.add_run(label+"  "),10.5,True,accent);set_run(p.add_run(text),10.2,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def add_image(doc,path,width=6.75):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(6)
    p.add_run().add_picture(str(path),width=Inches(width));return p


def add_lines(doc,count=3):
    for _ in range(count):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(5);set_run(p.add_run("________________________________________________________________________________"),9,False,LINE)


def add_page(doc):
    doc.add_page_break()


def simple_table(doc,headers,rows,widths,header_fill=NAVY,font_size=9.2):
    t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;set_table_widths(t,widths)
    for j,h in enumerate(headers):
        c=t.cell(0,j);set_cell_fill(c,header_fill);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_run(p.add_run(h),9,True,WHITE)
    for row in rows:
        cells=t.add_row().cells
        for j,val in enumerate(row):
            c=cells[j];c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(t.rows)%2==1:set_cell_fill(c,PALE)
            p=c.paragraphs[0];set_run(p.add_run(str(val)),font_size,False,INK)
    set_table_widths(t,widths);return t


def build_docx():
    doc=Document();style_doc(doc)
    # 1 cover
    sec=doc.sections[0];sec.header.is_linked_to_previous=False
    hp=sec.header.paragraphs[0];hp.clear();fp=sec.footer.paragraphs[0];fp.clear()
    add_image(doc,ASSET_DIR/"cover.png",width=7.06)
    add_page(doc)
    # restore furniture from page 2 onward by new section
    sec2=doc.add_section(WD_SECTION_START.NEW_PAGE);sec2.page_width=Inches(8.5);sec2.page_height=Inches(11);sec2.top_margin=sec2.bottom_margin=Inches(.72);sec2.left_margin=sec2.right_margin=Inches(.72);sec2.header_distance=Inches(.30);sec2.footer_distance=Inches(.32)
    sec2.header.is_linked_to_previous=False;sec2.footer.is_linked_to_previous=False
    hp=sec2.header.paragraphs[0];hp.alignment=WD_ALIGN_PARAGRAPH.LEFT;set_run(hp.add_run("CURIONEST CHEMISTRY  |  FOUNDATIONS PILOT"),8,True,GRAY)
    fp=sec2.footer.paragraphs[0];fp.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_run(fp.add_run("Single-classroom license  |  © 2026 CurioNest  |  Page "),8,False,GRAY);fld=OxmlElement("w:fldSimple");fld.set(qn("w:instr"),"PAGE");fp._p.append(fld)
    # 2 teacher guide
    add_title(doc,"Teacher Guide","A flexible 3-5 day launch unit for Grades 9-11","Start here")
    add_callout(doc,"PURPOSE","Students build safe lab habits, recognize essential equipment, and make defensible measurements before the first wet lab.")
    simple_table(doc,["Day","Focus","Suggested evidence"],[
        ("1","Safety norms + Safety Detective","Correct unsafe actions using evidence"),("2","GHS labels + SDS navigation","Locate hazard, PPE, first-aid information"),("3","Equipment identification + tool choice","Match tool to task and justify"),("4","Measurement stations","Read meniscus; tare and report mass"),("5","Practical + quiz","Demonstrate procedures and explain choices"),
    ],[720,3720,4920])
    doc.add_paragraph(style="Heading 2").add_run("Prep checklist")
    for text in ["Print pages labeled Student Page or Assessment; retain Teacher Page and Teacher Key pages.","Set out clean, empty equipment only. Do not use chemicals for the identification stations.","Test the eyewash and safety shower according to school policy; never activate them solely as a student demonstration.","Use local emergency procedures and the SDS for every chemical used in later labs."]:
        add_bullet(doc,text)
    add_callout(doc,"STANDARDS ROLE","This foundational skills unit supports NGSS science practices such as planning investigations, analyzing data, and using mathematics. It is not presented as a standalone HS-PS1 performance-expectation unit.")
    doc.add_paragraph(style="Heading 2").add_run("Suggested materials")
    add_para(doc,"Goggles, lab apron, beaker, Erlenmeyer flask, graduated cylinder, test tube and rack, dropper, funnel, electronic balance, wash bottle, ring stand, hot plate, and printed station cards.",10.2)
    add_callout(doc,"SAFETY NOTE","This resource supports instruction; it does not replace district policy, teacher training, equipment inspection, or chemical-specific SDS review.",fill="FFF4E8",accent=RED)
    # 3 safety readiness
    add_page(doc);add_title(doc,"Before You Touch the Tools","Student name: __________________  Class: __________  Date: __________","Student page 1")
    add_para(doc,"Complete the readiness check. Your teacher may add local procedures.",10.5,True,NAVY)
    rows=[
        ("Eye protection","Wear chemical splash goggles whenever chemicals, heat, or glassware are used as directed."),("Clothing","Closed-toe shoes; secure loose hair and clothing; wear the assigned apron."),("Behavior","Read the full procedure first. No food, drink, horseplay, or unauthorized experiments."),("Emergency","Know exits, gas/electric shutoffs, eyewash, shower, extinguisher, and spill-response procedures."),("Chemicals","Read the label and SDS. Never return unused chemical to the stock container."),("Incidents","Report every spill, breakage, exposure, injury, or near miss immediately."),
    ]
    simple_table(doc,["Readiness area","What safe practice looks like"],rows,[1900,7460],font_size=9.5)
    doc.add_paragraph(style="Heading 2").add_run("Explain, do not just agree")
    add_para(doc,"Why must splash goggles remain over the eyes rather than on the forehead?",10.2,True);add_lines(doc,2)
    add_para(doc,"Why should unused chemicals never be returned to the stock bottle?",10.2,True);add_lines(doc,2)
    add_para(doc,"Student commitment: I will stop and ask when a procedure, label, or safety direction is unclear.",10.2,True,TEAL_DARK)
    add_para(doc,"Student signature: __________________________   Date: __________",10.2)
    # 4 safety scene
    add_page(doc);add_title(doc,"Lab Safety Detective","Read eight classroom snapshots. Explain the safer action for each.","Student page 2")
    add_image(doc,ASSET_DIR/"safety-scene.png",width=6.9)
    simple_table(doc,["#","Unsafe choice","Safer action"],[(str(i),"","________________________________________") for i in range(1,9)],[620,4100,4640],font_size=9)
    # 5 GHS reference
    add_page(doc);add_title(doc,"GHS Pictogram Reference","Symbols communicate hazard classes; they do not replace the full label or SDS.","Student page 3")
    add_image(doc,ASSET_DIR/"ghs-grid.png",width=6.9)
    add_callout(doc,"REMEMBER","OSHA requires eight pictograms in its Hazard Communication Standard. The environmental pictogram is shown because students may see it, but OSHA does not mandate it.")
    # 6 label/SDS
    add_page(doc);add_title(doc,"Read a Chemical Label and SDS","Use the fictional acetone label below, then identify where the SDS supplies more detail.","Student page 4")
    add_image(doc,ASSET_DIR/"sample-label.png",width=6.9)
    qs=["1. Which pictograms appear, and what hazards do they communicate?","2. List two precautions stated on the label.","3. Which SDS section would you check for first-aid measures?","4. Which SDS section describes exposure controls and personal protection?","5. Why is a pictogram alone not enough information for safe use?"]
    for q in qs:add_para(doc,q,9.7,True,INK,after=2);add_lines(doc,1)
    add_callout(doc,"SDS QUICK MAP","1 Identification | 2 Hazard(s) | 4 First aid | 7 Handling/storage | 8 Exposure controls/PPE | 10 Stability/reactivity | 13 Disposal")
    # 7-8 equipment references
    for idx,file in enumerate(("equipment-reference-a.png","equipment-reference-b.png"),start=5):
        add_page(doc);add_title(doc,"Equipment Reference","Study the name and function. Circle tools available in your classroom.",f"Student page {idx}");add_image(doc,ASSET_DIR/file,width=6.9)
    # 9 identification
    add_page(doc);add_title(doc,"Equipment Identification Challenge","Write the correct name for each numbered tool.","Student page 7")
    add_image(doc,ASSET_DIR/"equipment-challenge.png",width=6.9)
    rows=[]
    for i in range(1,13,2):rows.append((f"{i}. __________________",f"{i+1}. __________________"))
    simple_table(doc,["Odd-numbered tools","Even-numbered tools"],rows,[4680,4680],font_size=10)
    # 10 tool choice
    add_page(doc);add_title(doc,"Choose the Right Tool","Name the best tool and justify why it fits the task.","Student page 8")
    tasks=[
        ("Measure 24.6 mL of water","",""),("Determine the mass of a metal sample","",""),("Add liquid one drop at a time","",""),("Heat a solution without an open flame","",""),("Support a buret during titration","",""),("Transfer a solid into a narrow-neck flask","",""),("Rinse the inside wall of glassware","",""),("Handle hot glassware safely","",""),
    ]
    simple_table(doc,["Task","Best tool","Reason"],tasks,[4000,2200,3160],font_size=9.3)
    doc.add_paragraph(style="Heading 2").add_run("Compare precision")
    add_para(doc,"Rank these from least to most appropriate for measuring 25.0 mL: beaker, graduated cylinder, buret. Explain your ranking.",10,True);add_lines(doc,3)
    # 11 measurement reference
    add_page(doc);add_title(doc,"Measurement Skills","Choose the instrument first; the number of reported digits follows the instrument.","Student page 9")
    add_image(doc,ASSET_DIR/"measurement-guide.png",width=6.9)
    add_callout(doc,"THREE HABITS","Read at eye level. Record every certain digit plus one estimated digit for an analog scale. Always include a unit.")
    # 12 meniscus practice
    add_page(doc);add_title(doc,"Meniscus and Precision Practice","Read the bottom of each meniscus at eye level.","Student page 10")
    add_image(doc,ASSET_DIR/"meniscus-practice.png",width=6.9)
    qs=["1. Which cylinder permits the most precise reading? Explain using the scale divisions.","2. A student records Cylinder A as 33 mL. What important measurement information is lost?","3. Describe one source of parallax error and how to prevent it."]
    for q in qs:add_para(doc,q,10,True,INK,after=2);add_lines(doc,2)
    # 13 station setup teacher
    add_page(doc);add_title(doc,"Lab Skills Circuit - Teacher Setup","Four low-risk stations using clean equipment and water only.","Teacher page")
    rows=[
        ("A: Identify","Display 8 tools","Students name each and state one safe use."),("B: Volume","Cylinder + colored water","Students read volume at eye level."),("C: Mass","Balance + empty container + sample","Students tare, measure, and report unit."),("D: Tool choice","Scenario cards","Students select a tool and justify precision/safety."),
    ]
    simple_table(doc,["Station","Setup","Performance evidence"],rows,[1450,3300,4610],font_size=9.3)
    doc.add_paragraph(style="Heading 2").add_run("Teacher moves")
    for x in ["Demonstrate one non-example before students rotate.","Require students to say both the value and unit aloud.","Do not place chemicals at these stations.","Use the practical rubric on page 15 for rapid feedback."]:
        add_bullet(doc,x)
    add_callout(doc,"DIFFERENTIATION","Provide the reference pages at the first rotation; remove them for the second. For advanced students, add a precision comparison between graduated and volumetric glassware.")
    # 14 station record
    add_page(doc);add_title(doc,"Lab Skills Circuit - Student Record","Name: ______________________________  Group: ______","Student page 11")
    simple_table(doc,["Station","Observation / measurement","Evidence or reason"],[
        ("A","Tool 1: __________  Tool 2: __________","Safe use: __________________________"),("B","Volume: __________ mL","Eye-level evidence: _________________"),("C","Container mass after tare: ______ g\nSample mass: ______ g","Why tare? _________________________"),("D","Scenario: ________________________\nChosen tool: ______________________","Why this tool? _____________________"),
    ],[1400,3900,4060],font_size=9.5)
    doc.add_paragraph(style="Heading 2").add_run("Exit reflection")
    add_para(doc,"One habit that prevents an inaccurate measurement is...",10,True);add_lines(doc,3)
    add_para(doc,"One safety action I can now explain—not merely memorize—is...",10,True);add_lines(doc,3)
    # 15 practical rubric
    add_page(doc);add_title(doc,"Lab Skills Practical","Teacher observation rubric; score each criterion 0-2.","Assessment")
    rubric=[
        ("PPE and readiness","Correct PPE; area clear; reads task first","Minor prompt needed","Unsafe/not ready"),("Tool identification","Names tool and appropriate use","Name or use incomplete","Incorrect"),("Volume technique","Eye level; bottom meniscus; unit","One technique error","Multiple errors"),("Balance technique","Tares; waits for stable reading; unit","One technique error","Multiple errors"),("Explanation","Uses safety/precision evidence","Partial reason","No defensible reason"),
    ]
    simple_table(doc,["Criterion","2 - Meets","1 - Developing","0 - Not yet"],rubric,[1750,2850,2500,2260],font_size=8.4)
    add_para(doc,"Score: ______ / 10       Teacher feedback: ______________________________________________",10.2,True,NAVY,after=8)
    add_lines(doc,4)
    add_callout(doc,"REASSESSMENT","Allow a second attempt after feedback. The goal is safe, repeatable performance—not a one-time score.")
    # 16 quiz 1
    add_page(doc);add_title(doc,"Foundations Quiz - Part A","Name: ______________________________  Class: ______","Assessment")
    mc=[
        ("1. Which tool is best for measuring 18.6 mL?",["Beaker","Graduated cylinder","Erlenmeyer flask","Watch glass"]),
        ("2. What should you do first after a chemical splashes into your eyes?",["Finish the procedure","Use the eyewash and alert the teacher","Wipe eyes with a towel","Look up the formula"]),
        ("3. Which pictogram indicates an oxidizer?",["Flame over circle","Gas cylinder","Exclamation mark","Skull and crossbones"]),
        ("4. Why is a balance tared?",["To heat the sample","To subtract the container mass","To change grams to mL","To increase sample mass"]),
        ("5. Which SDS section contains first-aid measures?",["Section 1","Section 2","Section 4","Section 10"]),
        ("6. Which action is safest?",["Smell directly from a bottle","Return excess chemical to stock","Tie back loose hair","Wear goggles on the forehead"]),
    ]
    for q,choices in mc:
        add_para(doc,q,9.8,True,NAVY,after=1)
        add_para(doc,"    ".join(f"{chr(65+i)}. {c}" for i,c in enumerate(choices)),9.2,after=6)
    doc.add_paragraph(style="Heading 2").add_run("Match the tool to the task")
    simple_table(doc,["Task","Tool"],[("7. Add liquid dropwise","______________"),("8. Support a buret","______________"),("9. Heat without an open flame","______________"),("10. Measure mass","______________")],[5600,3760],font_size=9.5)
    # 17 quiz part B
    add_page(doc);add_title(doc,"Foundations Quiz - Part B","Explain your reasoning using safety or measurement evidence.","Assessment")
    prompts=[
        "11. A student reads a graduated cylinder while looking down from above. Identify the likely error and describe the correction.",
        "12. The label shows a flame pictogram and the word DANGER. State two precautions that should be checked before use.",
        "13. A classmate wants to measure 25.0 mL with a beaker because it is faster. Write a claim, evidence, and reasoning response.",
        "14. During cleanup, you find broken glass in the sink. Explain the safe response.",
    ]
    for p in prompts:
        add_para(doc,p,10,True,NAVY,after=2);add_lines(doc,4)
    # 18 key safety/GHS
    add_page(doc);add_title(doc,"Answer Key - Safety and GHS","Suggested responses; accept equivalent scientifically sound wording.","Teacher key 1")
    key=[
        ("Safety Detective 1-4","1 goggles over eyes; 2 remove drink; 3 label container; 4 move paper/flammables away and control flame."),("Safety Detective 5-8","5 notify teacher; use broken-glass bin with proper tools; 6 remove backpack from aisle; 7 do not leave a reagent container open or unlabeled; 8 secure loose hair."),("Label Q1","Flame = flammable hazard; exclamation = irritant/harmful effects."),("Label Q2","Any two: keep away from ignition; wear eye protection; use adequate ventilation."),("Label Q3-4","First aid = SDS Section 4. Exposure controls/PPE = Section 8."),("Label Q5","The label is a quick warning; the SDS provides detailed controls, first aid, storage, stability, and disposal information."),
    ]
    simple_table(doc,["Item","Suggested answer"],key,[2200,7160],font_size=9.2)
    doc.add_paragraph(style="Heading 2").add_run("Readiness explanations")
    add_para(doc,"Goggles must cover the eyes to protect against splashes. Chemicals are not returned to stock because they may contaminate the original reagent.",10)
    # 19 key equipment
    add_page(doc);add_title(doc,"Answer Key - Equipment and Measurement","Use the reference pages to reteach before reassessment.","Teacher key 2")
    names=[name for name,_ in EQUIPMENT[:12]]
    simple_table(doc,["#","Equipment","#","Equipment"],[(i+1,names[i],i+7,names[i+6]) for i in range(6)],[650,4030,650,4030],font_size=9.3)
    doc.add_paragraph(style="Heading 2").add_run("Choose the Right Tool")
    answers=[("24.6 mL","Graduated cylinder","Graduations support a measured volume."),("Mass","Electronic balance","Tare container; report g."),("Dropwise","Dropper pipet","Controls small additions."),("Heat/no flame","Hot plate","No open flame."),("Support buret","Ring stand","Holds a clamp and buret upright."),("Solid to narrow flask","Funnel","Guides transfer."),("Rinse glassware","Wash bottle","Directed stream."),("Hot glassware","Crucible tongs","Keep hands away from hot surfaces.")]
    simple_table(doc,["Task","Answer","Reason"],answers,[2500,2500,4360],font_size=8.8)
    add_para(doc,"Precision ranking: beaker < graduated cylinder < buret for variable-volume measurement. Suitability still depends on the specific task and required uncertainty.",9.5,True,TEAL_DARK)
    # 20 key quiz/references
    add_page(doc);add_title(doc,"Answer Key - Measurement, Quiz and Sources","Suggested answers and scoring notes for consistent feedback.","Teacher key 3")
    add_para(doc,"Meniscus practice: A = 32.6 mL; B = 18.4 mL; C = 46.5 mL. Cylinder A/B have 1 mL minor divisions, so readings should include an estimated tenths place. Cylinder C has 2 mL divisions and supports an estimated reading between marks. Parallax occurs when the eye is above or below the meniscus; move to eye level.",9.6)
    simple_table(doc,["Quiz","Answer / scoring note"],[("1-6","B, B, A, B, C, C"),("7-10","Dropper pipet; ring stand; hot plate; electronic balance"),("11","Viewing from above causes parallax; place eyes level with the meniscus and read the bottom."),("12","Any two defensible label/SDS precautions, such as remove ignition sources, wear goggles, or provide ventilation."),("13","Claim: use a graduated cylinder. Evidence: its scale supports a more precise reading than a beaker. Reasoning links smaller divisions to lower reading uncertainty."),("14","Notify the teacher; do not pick up by hand; use designated tools and place glass in the broken-glass container per local procedure.")],[1300,8060],font_size=9.0)
    doc.add_paragraph(style="Heading 2").add_run("Teacher safety references")
    add_para(doc,"American Chemical Society. Student Laboratory Code of Conduct; Lab & Safety Equipment; Importance of the Laboratory Experience. OSHA. Hazard Communication Standard / GHS Questions and Answers.",8.8,False,GRAY)
    add_para(doc,"https://institute.acs.org/acs-center/lab-safety/education-training/high-school-labs/student-lab-code-of-conduct.html",8.0,False,BLUE)
    add_para(doc,"https://www.acs.org/education/policies/middle-and-high-school-chemistry/classroom-and-lab-facilities/safety-equipment.html",8.0,False,BLUE)
    add_para(doc,"https://www.osha.gov/hazcom/faq",8.0,False,BLUE)
    doc.add_paragraph(style="Heading 2").add_run("Visual asset credits")
    add_para(doc,"Equipment icons by Servier Medical Art (CC BY 3.0) and OpenClipart contributors (CC0), accessed through Bioicons. GHS pictograms by UNECE (CC0), accessed through Bioicons. Individual file URLs and licenses are recorded in assets-manifest.json.",7.8,False,GRAY)
    add_callout(doc,"TERMS OF USE","Single-classroom use by the original purchaser. Do not post publicly, resell, or redistribute editable files. School/district licenses require additional permission.",fill=PALE,accent=NAVY)
    OUT.mkdir(parents=True,exist_ok=True);doc.save(DOCX_PATH);return DOCX_PATH


def main():
    prepare_assets();path=build_docx();print(path)


if __name__=="__main__":
    main()
