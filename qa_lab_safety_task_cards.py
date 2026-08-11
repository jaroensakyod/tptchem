"""Deterministic package checks for Lab Safety Task Cards."""

from pathlib import Path
import hashlib
import json
import re
import sys
import zipfile

from docx import Document
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
PKG = ROOT / "products" / "lab-safety-task-cards"
SOURCE = PKG / "source" / "source.json"
BUYER = PKG / "output" / "buyer"
STATIONS = PKG / "output" / "stations"
LISTING = PKG / "output" / "listing"

errors = []


def check(condition, message):
    if not condition:
        errors.append(message)


def norm(value):
    return " ".join(str(value).split()).lower()


data = json.loads(SOURCE.read_text(encoding="utf-8"))
cards = data.get("cards", [])
cards_hash = hashlib.sha256(json.dumps(cards, sort_keys=True).encode("utf-8")).hexdigest()
check(data.get("review", {}).get("cards_sha256") == cards_hash, "card review hash is stale")
check(len(cards) == 24, "source must contain exactly 24 cards")
check([card.get("id") for card in cards] == [f"T{i:02d}" for i in range(1, 25)], "card IDs must be T01-T24 in order")
check(len({card.get("title") for card in cards}) == 24, "card titles must be unique")

required = ("portion", "portion_name", "level", "type", "title", "prompt", "options", "answer", "evidence", "risk", "action", "why", "acceptable", "misconception", "sources")
for card in cards:
    for field in required:
        check(bool(card.get(field)), f"{card.get('id')} missing {field}")
    check(len(card.get("options", [])) == 4, f"{card.get('id')} must contain four options")
    answers = re.split(r"\s+and\s+|\s*,\s*", card.get("answer", ""))
    check(bool(answers) and all(answer in "ABCD" for answer in answers), f"{card.get('id')} has an invalid answer key")

check({portion: sum(card.get("portion") == portion for card in cards) for portion in "ABCD"} == {portion: 6 for portion in "ABCD"}, "each portion must contain six cards")
check({level: sum(card.get("level") == level for card in cards) for level in ("Support", "Core", "Honors")} == {"Support": 8, "Core": 12, "Honors": 4}, "level counts do not match the product brief")

visual_cards = [card for card in cards if card.get("visual")]
visual_gate = data.get("visual_gate", {})
required_visual_ids = visual_gate.get("required_visual_ids", [])
completed_visual_ids = visual_gate.get("completed_visual_ids", [])
actual_visual_ids = [card["id"] for card in visual_cards]
check(completed_visual_ids == actual_visual_ids, "visual gate completed IDs do not match actual card placements")
check(visual_gate.get("required_count") == len(required_visual_ids), "visual gate required_count is inconsistent")
check(visual_gate.get("completed_count") == len(actual_visual_ids), "visual gate completed_count is inconsistent")
check(visual_gate.get("missing_count") == len(set(required_visual_ids) - set(actual_visual_ids)), "visual gate missing_count is inconsistent")
check(
    visual_gate.get("status") == "passed" and set(required_visual_ids).issubset(actual_visual_ids),
    f"instructional visual gate BLOCKED: {len(actual_visual_ids)}/{len(required_visual_ids)} required card visuals complete",
)
for card in visual_cards:
    asset = ROOT / card["visual"]["asset"]
    check(asset.is_file(), f"{card['id']} visual asset is missing: {card['visual']['asset']}")
check(len(data.get("visual_sources", [])) == 3, "visual source ledger must contain three source records")
for source in data.get("visual_sources", []):
    for field in ("name", "creator", "url", "license", "use"):
        check(bool(source.get(field)), f"visual source missing {field}")
check(not list((PKG / "source" / "visuals").glob("*.png")), "source/visuals must not contain generated or code-drawn PNG assets")

pdfs = {
    BUYER / "product.pdf": 30,
    BUYER / "student-packet.pdf": 15,
    BUYER / "student-packet-bw.pdf": 15,
    BUYER / "teacher-guide-and-key.pdf": 16,
    LISTING / "preview.pdf": 3,
}
for path, expected_pages in pdfs.items():
    check(path.is_file() and path.stat().st_size > 1000, f"{path.name} missing or too small")
    if path.is_file():
        reader = PdfReader(str(path))
        check(len(reader.pages) == expected_pages, f"{path.name} expected {expected_pages} pages")
        for page_number, page in enumerate(reader.pages, 1):
            check(abs(float(page.mediabox.width) - 612) < 1 and abs(float(page.mediabox.height) - 792) < 1, f"{path.name} page {page_number} is not US Letter")

product_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(BUYER / "product.pdf")).pages)
for card in cards:
    check(norm(card["title"]) in norm(product_text), f"product missing {card['id']} title")
    for field in ("evidence", "risk", "action", "why", "misconception"):
        check(norm(card[field]) in norm(product_text), f"product key missing {field} for {card['id']}")

student_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(BUYER / "student-packet.pdf")).pages)
check("CORRECT ANSWER" not in student_text.upper() and "ANSWER KEY" not in student_text.upper(), "student packet exposes answer content")
teacher_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(BUYER / "teacher-guide-and-key.pdf")).pages)
check("ANSWER KEY" in teacher_text.upper() and "MISCONCEPTION" in teacher_text.upper(), "teacher guide is missing the full key")

portion_names = {
    "A": "portion-a-notice-and-pause.pdf",
    "B": "portion-b-choose-next-move.pdf",
    "C": "portion-c-communicate-procedure.pdf",
    "D": "portion-d-transfer-repair.pdf",
}
for portion, filename in portion_names.items():
    path = STATIONS / filename
    check(path.is_file(), f"{filename} is missing")
    if path.is_file():
        reader = PdfReader(str(path))
        check(len(reader.pages) == 8, f"{filename} must contain 8 pages")
        text = norm("\n".join((page.extract_text() or "") for page in reader.pages))
        for card in cards:
            if card["portion"] == portion:
                check(norm(card["title"]) in text, f"{filename} missing {card['id']}")

docx = BUYER / "product-editable.docx"
check(docx.is_file(), "editable DOCX is missing")
if docx.is_file():
    document = Document(str(docx))
    doc_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    doc_text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    for card in cards:
        check(norm(card["title"]) in norm(doc_text), f"editable DOCX missing {card['id']}")
    with zipfile.ZipFile(docx) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    check(len(re.findall(r'\bdescr="[^"]+"', xml)) == 4, "editable DOCX must contain four non-empty image descriptions")

for name, dimensions in {
    "cover.png": (1200, 1200),
    "listing-01-cover.png": (1800, 1800),
    "listing-02-inside.png": (1800, 1800),
    "listing-03-teacher-ready.png": (1800, 1800),
}.items():
    path = LISTING / name
    check(path.is_file(), f"{name} is missing")
    if path.is_file():
        with Image.open(path) as image:
            check(image.size == dimensions, f"{name} has incorrect dimensions")

root_files = sorted(path.name for path in PKG.iterdir() if path.is_file())
check(root_files == ["README.md"], f"package root contains misplaced files: {root_files}")

if errors:
    for message in errors:
        print("FAIL", message)
    print(f"SUMMARY: {len(errors)} failure(s)")
    sys.exit(1)

print("PASS Lab Safety Task Cards deterministic package QA")
print(f"SUMMARY: 24 cards, {len(actual_visual_ids)} verified visual placements, 9 PDFs, 1 editable DOCX, 0 failures")
