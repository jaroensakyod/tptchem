"""Deterministic package checks for Lab Safety Scenario Analysis."""
from pathlib import Path
import json
import hashlib
import sys

from docx import Document
from PIL import Image
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent
PKG = ROOT / "products" / "lab-safety-scenario-analysis"
SOURCE = PKG / "source" / "source.json"
BUYER = PKG / "output" / "buyer"
LISTING = PKG / "output" / "listing"

errors=[]
def check(condition, message):
    if not condition: errors.append(message)

def norm(value): return " ".join(str(value).split()).lower()

source=json.loads(SOURCE.read_text(encoding="utf-8"))
items=source.get("scenarios",[])
scenario_hash=hashlib.sha256(json.dumps(items,sort_keys=True,separators=(",",":"),ensure_ascii=False).encode("utf-8")).hexdigest()
check(source.get("review",{}).get("scenario_content_sha256")==scenario_hash,"scenario review hash is stale")
check(source.get("review",{}).get("safety_status")=="internal_qa_approved_pending_classroom_review","safety review status missing")
check(len(items)==10,"source must contain exactly 10 scenarios")
ids=[x.get("id") for x in items]
check(ids==list(range(1,11)),"scenario IDs must be 1-10 in order")
check(len({x.get("title") for x in items})==10,"scenario titles must be unique")
for x in items:
    for field in ("title","text","evidence","risk","action","why"):
        check(bool(str(x.get(field,"")).strip()),f"scenario {x.get('id')} missing {field}")

visual_gate=source.get("visual_gate",{})
required_visual_ids=visual_gate.get("required_visual_ids",[])
actual_visual_ids=[x.get("id") for x in items if x.get("visual")]
check(visual_gate.get("completed_visual_ids",[])==actual_visual_ids,"visual gate completed IDs do not match actual scenario placements")
check(visual_gate.get("required_count")==len(required_visual_ids),"visual gate required_count is inconsistent")
check(visual_gate.get("completed_count")==len(actual_visual_ids),"visual gate completed_count is inconsistent")
check(visual_gate.get("missing_count")==len(set(required_visual_ids)-set(actual_visual_ids)),"visual gate missing_count is inconsistent")
check(
    visual_gate.get("status")=="passed" and set(required_visual_ids).issubset(actual_visual_ids),
    f"instructional visual gate BLOCKED: {len(actual_visual_ids)}/{len(required_visual_ids)} required scenario visuals complete",
)

expected={"product.pdf":12,"student-packet.pdf":6,"student-packet-bw.pdf":6,"teacher-guide-and-key.pdf":6,"preview.pdf":3}
for name,count in expected.items():
    path=(LISTING if name=="preview.pdf" else BUYER)/name; check(path.exists() and path.stat().st_size>1000,f"{name} missing or too small")
    if path.exists():
        reader=PdfReader(str(path)); check(len(reader.pages)==count,f"{name} expected {count} pages")
        for n,page in enumerate(reader.pages,1):
            w=float(page.mediabox.width); h=float(page.mediabox.height)
            check(abs(w-612)<1 and abs(h-792)<1,f"{name} page {n} is not US Letter")
            if name!="student-packet-bw.pdf":
                check(len((page.extract_text() or "").strip())>80,f"{name} page {n} appears blank")

product_text="\n".join((p.extract_text() or "") for p in PdfReader(str(BUYER/"product.pdf")).pages)
for x in items:
    check(norm(x["title"]) in norm(product_text),f"product missing scenario title {x['id']}")
    for field in ("evidence","risk","action","why"):
        check(norm(x[field]) in norm(product_text),f"product key missing {field} {x['id']}")
check("contaminated glass" not in norm(product_text),"product invents glass contamination")
check("□" not in product_text and "■" not in product_text,"product contains broken checkbox glyph")
student_text="\n".join((p.extract_text() or "") for p in PdfReader(str(BUYER/"student-packet.pdf")).pages)
check("ANSWER KEY" not in student_text.upper(),"student packet exposes answer key")
check("STUDENT 1 OF 6" in student_text.upper() and "STUDENT 6 OF 6" in student_text.upper(),"student packet footer is not repaginated")
teacher_text="\n".join((p.extract_text() or "") for p in PdfReader(str(BUYER/"teacher-guide-and-key.pdf")).pages)
check("ANSWER KEY" in teacher_text.upper() and "CER MODEL" in teacher_text.upper(),"teacher file missing full key or CER model")
preview_text="\n".join((p.extract_text() or "") for p in PdfReader(str(LISTING/"preview.pdf")).pages)
check(preview_text.upper().count("PREVIEW")>=expected["preview.pdf"],"preview watermark missing from one or more pages")

# ReportLab retains an unused Helvetica resource for empty BT/ET setup commands.
# Require every font that actually draws text to be the embedded OFL Poppins font.
reader=PdfReader(str(BUYER/"product.pdf")); embedded_poppins=0
for page in reader.pages:
    fonts=(page["/Resources"].get("/Font") or {}).get_object()
    for font_ref in fonts.values():
        font=font_ref.get_object(); base=str(font.get("/BaseFont",""))
        descriptor=font.get("/FontDescriptor"); embedded=False
        if descriptor:
            desc=descriptor.get_object(); embedded=any(k in desc for k in ("/FontFile","/FontFile2","/FontFile3"))
        if "Poppins" in base:
            embedded_poppins+=1; check(embedded,"rendered Poppins font is not embedded")
    content=page.get_contents().get_data()
    for segment in content.split(b"BT"):
        if b"/F1 " in segment and (b" Tj" in segment or b" TJ" in segment):
            errors.append("unembedded Helvetica draws visible text")
check(embedded_poppins>0,"embedded Poppins fonts not found")

doc=Document(str(BUYER/"product-editable.docx"))
doc_text="\n".join(p.text for p in doc.paragraphs)+"\n"+"\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
check("lab safety" in norm(doc_text) and "sources, rights, and use" in norm(doc_text),"editable DOCX missing required sections")
for x in items: check(norm(x["title"]) in norm(doc_text),f"editable DOCX missing scenario {x['id']}")
check(doc.core_properties.author=="CurioNest","editable DOCX author metadata not scrubbed")
check(doc.core_properties.title==source["title"],"editable DOCX title metadata mismatch")

for name in ("cover.png","listing-01-cover.png","listing-02-inside.png","listing-03-teacher-ready.png"):
    path=LISTING/name; check(path.exists(),f"{name} missing")
    if path.exists():
        expected_size=(1200,1200) if name=="cover.png" else (1800,1800)
        check(Image.open(path).size==expected_size,f"{name} must be {expected_size[0]}x{expected_size[1]}")

if errors:
    print("FAIL")
    for e in errors: print("-",e)
    sys.exit(1)
print("PASS: source, keys, DOCX, PDFs, page sizes, student split, preview watermark, and listing images")
