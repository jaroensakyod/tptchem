#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""CurioNest Quiz builder — standalone quiz (MCQ / T-F / calculation) → docx → PDF.

Schema (mirror-json/quiz-*.json):
{
  "title": "...", "subtitle": "...", "meta": {"file_stem": "...", "duration": "20 minutes"},
  "sections": [
    {"kind":"header"},
    {"kind":"instructions", "items": ["..."], "points": 20},
    {"kind":"quiz", "title":"Part A — Multiple Choice", "items":[
        {"q":"...", "type":"mc", "choices":["A","B","C","D"], "answer":"B", "explain":"..."},
        {"q":"...", "type":"tf", "answer":"True", "explain":"..."},
        {"q":"...", "type":"calc", "answer":"0.025 M", "points":4, "explain":"..."}
    ]},
    {"kind":"answer_key"}   # auto-built from quiz items
  ]
}

Same proven pipeline: font vertAlign for inline sub/sup, $...$ OMML for display,
US Letter, Georgia, page border, header/footer, Word COM → PDF.
"""
import json, os, re, sys

BASE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INDIGO = RGBColor(0x31, 0x2E, 0x81)
PURPLE = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x33, 0x41, 0x55)
GRAY = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0x9F, 0x12, 0x39)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x15, 0x80, 0x3D)

def add_chem_text(par, text, size=10.5, bold=False, color=None, italic=False):
    from lxml import etree
    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    for seg in re.split(r"(\$[^$]+\$)", text):
        if not seg:
            continue
        if seg.startswith("$") and seg.endswith("$") and len(seg) > 2:
            try:
                from omml import latex_to_omml
                body = latex_to_omml(seg[1:-1])
                omath = etree.fromstring(
                    f'<m:oMath xmlns:m="{M_NS}">'
                    f'<m:r><m:rPr><m:sz m:val="{int(size * 2)}"/></m:rPr></m:r>{body}'
                    f"</m:oMath>")
                par._p.append(omath)
            except Exception:
                r = par.add_run(seg)
                r.font.size = Pt(size); r.font.bold = bold
                if color: r.font.color.rgb = color
            continue
        stack = []
        pos = 0
        for m in re.finditer(r"<(/?)(sub|sup|b|i|br)([^>]*)>", seg):
            if m.start() > pos:
                run = par.add_run(seg[pos:m.start()])
                run.font.size = Pt(size)
                run.font.bold = bold or "b" in stack
                run.font.italic = italic or "i" in stack
                if color: run.font.color.rgb = color
                if "sub" in stack: run.font.subscript = True
                if "sup" in stack: run.font.superscript = True
            closing, tag, _ = m.groups()
            if tag == "br":
                par.add_run().add_break()
            elif closing:
                if tag in stack: stack.remove(tag)
            else:
                stack.append(tag)
            pos = m.end()
        if pos < len(seg):
            run = par.add_run(seg[pos:])
            run.font.size = Pt(size)
            run.font.bold = bold or "b" in stack
            run.font.italic = italic or "i" in stack
            if color: run.font.color.rgb = color
            if "sub" in stack: run.font.subscript = True
            if "sup" in stack: run.font.superscript = True

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcPr.append(shd)

def cell_borders(cell, color="999999", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.65)
    sec.top_margin, sec.bottom_margin = Inches(0.55), Inches(0.6)
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(2)
    return doc

def add_page_border(doc, color="4F46E5", sz="10"):
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "18")
        el.set(qn("w:color"), color)
        pgBorders.append(el)
    sectPr.append(pgBorders)

def add_header_footer(doc, brand_text):
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(brand_text)
    r.font.size = Pt(7.5); r.font.color.rgb = GRAY; r.font.name = "Georgia"
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = fp.add_run("© 2026 CurioNest · For classroom use only · Page ")
    r2.font.size = Pt(7.5); r2.font.color.rgb = GRAY
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    fp.add_run()._r.append(fld1)
    fp.add_run()._r.append(instr)
    fp.add_run()._r.append(fld2)

def para(doc, text, size=10, color=SLATE, bold=False, italic=False, after=3,
         align=WD_ALIGN_PARAGRAPH.LEFT, indent=None, before=0, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if before: p.paragraph_format.space_before = Pt(before)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if keep: p.paragraph_format.keep_with_next = True
    add_chem_text(p, text, size=size, bold=bold, color=color, italic=italic)
    return p

def heading(doc, text, size=12, color=INDIGO, bold=True, after=4, before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    add_chem_text(p, text, size=size, bold=bold, color=color)
    return p

# ---------------------------------------------------------------- renderers
def render_header(doc, s):
    para(doc, "", size=4, after=0)
    para(doc, "CURIONEST", size=13, color=INDIGO, bold=True, after=4,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, s.get("title", ""), size=19, color=INDIGO, bold=True, after=3,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, s.get("subtitle", ""), size=10, color=GRAY, italic=True, after=12,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    # quiz info table
    meta = s.get("meta", {})
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [("Duration", meta.get("duration", "20 minutes")),
            ("Questions", str(meta.get("n_questions", "—"))),
            ("Points", str(meta.get("points", "—"))),
            ("Answer Key", "Included — full explanations")]
    for i, (k, v) in enumerate(rows):
        c0, c1 = tbl.cell(i, 0), tbl.cell(i, 1)
        shade_cell(c0, "EEF2FF")
        add_chem_text(c0.paragraphs[0], k, size=9, bold=True)
        add_chem_text(c1.paragraphs[0], v, size=9, color=SLATE)
    para(doc, "Made by a chemistry teacher — aligned to the standard US chemistry sequence.",
         size=8.5, color=GRAY, italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

def render_instructions(doc, s):
    heading(doc, "Instructions", size=11.5, before=0)
    for it in s.get("items", []):
        para(doc, "•  " + it, size=9.8, indent=0.2, after=2)
    pts = s.get("points")
    if pts:
        para(doc, f"Total: {pts} points", size=9.8, color=RED, bold=True, after=6)

def render_quiz(doc, s, num_start):
    heading(doc, s.get("title", "Quiz"), size=11.5, before=8)
    n = num_start
    for item in s.get("items", []):
        t = item.get("type", "mc")
        pts = item.get("points", 1)
        # question line with points
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        add_chem_text(p, f"{n}.  ({pts} pt)  " if pts != 1 else f"{n}.  ",
                      size=9.8, bold=True, color=INDIGO)
        add_chem_text(p, item.get("q", ""), size=9.8)
        if t == "mc":
            for j, ch in enumerate(item.get("choices", [])):
                para(doc, f"    ({'ABCD'[j]})  {ch}", size=9, color=SLATE, after=0, keep=True)
        elif t == "tf":
            para(doc, "    True / False", size=9, color=GRAY, italic=True, after=0, keep=True)
        # answer space: MCQ = circle letter, calc = box, short = lines
        if t == "mc":
            para(doc, "    Answer: ______", size=9, color=GRAY, after=5)
        elif t == "tf":
            para(doc, "    Answer: ______", size=9, color=GRAY, after=5)
        else:  # calc
            for _ in range(2):
                para(doc, "", size=8, after=6)
        n += 1
    return n

def render_answer_key(doc, s, quiz_sections):
    doc.add_page_break()
    heading(doc, "Answer Key", size=13, before=0)
    para(doc, "Full explanations for every question — great for review or sub days.",
         size=9.5, color=GRAY, italic=True, after=8)
    for qs in quiz_sections:
        heading(doc, qs.get("title", ""), size=10.5, before=6)
        for i, item in enumerate(qs.get("items", [])):
            ans = item.get("answer", "")
            if not ans:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{qs['_start'] + i}.  ")
            r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INDIGO
            add_chem_text(p, ans, size=9.5, color=GREEN, bold=True)
            if item.get("explain"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(3)
                p2.paragraph_format.left_indent = Inches(0.3)
                r2 = p2.add_run("Why: ")
                r2.font.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = GRAY
                add_chem_text(p2, item["explain"], size=9, color=GRAY)

# ---------------------------------------------------------------- main
def build(data):
    doc = new_doc()
    add_page_border(doc)
    add_header_footer(doc, "CurioNest — " + data.get("title", ""))
    quiz_sections = [s for s in data.get("sections", []) if s.get("kind") == "quiz"]
    n = 1
    for qs in quiz_sections:
        qs["_start"] = n
        n += len(qs.get("items", []))
    meta = data.get("meta", {})
    meta["n_questions"] = n - 1
    for s in data.get("sections", []):
        kind = s.get("kind")
        if kind == "header":
            s["meta"] = meta
            render_header(doc, s)
        elif kind == "instructions":
            render_instructions(doc, s)
        elif kind == "quiz":
            n = render_quiz(doc, s, s["_start"])
        elif kind == "answer_key":
            render_answer_key(doc, s, quiz_sections)
        else:
            print(f"  (skip unknown kind: {kind})")
    return doc

def main():
    if len(sys.argv) < 2:
        sys.exit("usage: python build_quiz.py mirror-json/quiz-x.json [outstem]")
    src = sys.argv[1]
    data = json.load(open(src, encoding="utf-8"))
    stem = sys.argv[2] if len(sys.argv) > 2 else data.get("meta", {}).get("file_stem", "quiz")
    docx_path = os.path.join(BASE, f"{stem}-docx.docx")
    build(data).save(docx_path)
    print(f"docx saved: {docx_path}")
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(docx_path))
        pdf_path = os.path.join(BASE, f"{stem}-docx.pdf")
        d.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        d.Close(False)
        word.Quit()
        print(f"pdf saved: {pdf_path}")
    except Exception as e:
        print(f"PDF conversion failed: {e}")

if __name__ == "__main__":
    main()
