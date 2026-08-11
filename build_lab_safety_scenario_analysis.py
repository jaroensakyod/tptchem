from pathlib import Path
import json
from datetime import datetime
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "products" / "lab-safety-scenario-analysis"
DATA = json.loads((OUT / "source.json").read_text(encoding="utf-8"))
DOCX = OUT / "product-editable.docx"

NAVY = "102A43"; TEAL = "007C83"; PALE = "EAF7F7"; GOLD = "F2C14E"; RED = "C44536"; GRAY = "E8EEF2"; WHITE = "FFFFFF"

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), color)

def margins(section):
    section.top_margin = Inches(.55); section.bottom_margin = Inches(.55)
    section.left_margin = Inches(.62); section.right_margin = Inches(.62)

def footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Lab Safety Scenario Analysis  |  CurioNest  |  ")
    r.font.name = "Aptos"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string("66788A")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

def text(p, s, size=10, bold=False, color=NAVY):
    r = p.add_run(s); r.font.name = "Aptos"; r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = RGBColor.from_string(color); return r

def title(doc, kicker, heading, sub=""):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); text(p, kicker.upper(), 9, True, TEAL)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5); text(p, heading, 22, True, NAVY)
    if sub:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); text(p, sub, 10, False, "435B6B")

def banner(doc, s, color=TEAL):
    t=doc.add_table(rows=1, cols=1); t.autofit=False; t.columns[0].width=Inches(7.15)
    c=t.cell(0,0); shade(c,color); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(p,s,11,True,WHITE)

def box(doc, head, body, fill=PALE):
    t=doc.add_table(rows=2, cols=1); t.autofit=False; t.columns[0].width=Inches(7.15)
    shade(t.cell(0,0), TEAL); text(t.cell(0,0).paragraphs[0], head, 10, True, WHITE)
    shade(t.cell(1,0), fill); p=t.cell(1,0).paragraphs[0]; p.paragraph_format.space_after=Pt(3); text(p, body, 9)

def page(doc):
    # A next-page section is stable after large tables. A standalone page-break
    # paragraph can be pushed to a new page by Word and create a blank page.
    doc.add_section(WD_SECTION.NEW_PAGE)

def add_lines(cell, n=3):
    for _ in range(n):
        p=cell.add_paragraph("____________________________________________________________________")
        p.paragraph_format.space_after=Pt(2); text(p,"",8)

def scenario_card(doc, item, compact=False):
    t=doc.add_table(rows=2, cols=1); t.autofit=False; t.columns[0].width=Inches(7.15)
    shade(t.cell(0,0), NAVY); p=t.cell(0,0).paragraphs[0]; text(p,f"{item['id']}. {item['title']}",10,True,WHITE)
    c=t.cell(1,0); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); text(p,item['text'],9)
    if not compact:
        for label in ("Evidence noticed","Risk / possible harm","Safest first action","Why this action reduces risk"):
            p=c.add_paragraph(); p.paragraph_format.space_after=Pt(1)
            text(p,label+": ",8,True,TEAL); text(p,"____________________________________________",8,False,"66788A")
            if label.startswith("Why"):
                p2=c.add_paragraph("____________________________________________________________________")
                p2.paragraph_format.space_after=Pt(1)
    # Avoid an empty trailing paragraph: after five cards it can force the
    # following section break onto an otherwise blank page in Microsoft Word.

def answer_card(doc, item):
    t=doc.add_table(rows=2,cols=1); t.autofit=False; t.columns[0].width=Inches(7.15)
    shade(t.cell(0,0),NAVY); text(t.cell(0,0).paragraphs[0],f"{item['id']}. {item['title']}",9,True,WHITE)
    c=t.cell(1,0)
    fields=[("Evidence",item["evidence"]),("Risk",item["risk"]),("Action",item["action"]),("Why",item["why"])]
    for i,(label,value) in enumerate(fields):
        p=c.paragraphs[0] if i==0 else c.add_paragraph()
        p.paragraph_format.space_after=Pt(1); text(p,label+": ",7.5,True,TEAL); text(p,value,7.5)

doc=Document(); margins(doc.sections[0]); footer(doc.sections[0])
doc.core_properties.title=DATA["title"]
doc.core_properties.author="CurioNest"
doc.core_properties.subject="Lab safety reasoning for grades 9-11 chemistry"
doc.core_properties.keywords="lab safety, chemistry, scenario analysis, grades 9-11"
doc.core_properties.comments="Original classroom resource generated from products/lab-safety-scenario-analysis/source.json"
doc.core_properties.created=datetime(2026,8,11)
doc.core_properties.modified=datetime(2026,8,11)
styles=doc.styles
styles["Normal"].font.name="Aptos"; styles["Normal"].font.size=Pt(10); styles["Normal"].font.color.rgb=RGBColor.from_string(NAVY)

# 1 cover
banner(doc,"CHEMISTRY FOUNDATIONS • FORMAT 1")
doc.add_paragraph().paragraph_format.space_after=Pt(20)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(p,"LAB SAFETY",28,True,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(p,"SCENARIO ANALYSIS",28,True,TEAL)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(p,"MAKE THE SAFE CALL",14,True,RED)
doc.add_paragraph().paragraph_format.space_after=Pt(18)
box(doc,"WHAT STUDENTS DO","Analyze 10 original laboratory situations, identify evidence of risk, choose the safest first action, and justify decisions using a consistent safety framework.")
doc.add_paragraph()
t=doc.add_table(rows=2,cols=3); t.style="Table Grid"
for i,s in enumerate(["GRADES 9-11","CORE 45-55 MIN","PRINT + EDITABLE","CORE + HONORS","EXIT TICKET","FULL KEY"]):
    c=t.cell(i//3,i%3); shade(c, PALE if i<3 else GRAY); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(p,s,9,True,TEAL)
doc.add_paragraph().paragraph_format.space_after=Pt(26)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; text(p,"Original content • Printer-friendly • Structured reasoning",10,True,NAVY)

# 2 welcome / teacher guide
page(doc); title(doc,"Welcome / Teacher Guide","Welcome, Teacher","Use this ready-to-teach lesson for direct instruction, collaborative practice, a sub plan, or a readiness check.")
box(doc,"LEARNING TARGET","I can notice evidence, describe the risk, choose the safest immediate action, and explain why it reduces harm.")
doc.add_paragraph(); banner(doc,"SUGGESTED 50-MINUTE FLOW",NAVY)
flow=[("5 min","Launch: What makes an action the safest first move?"),("8 min","Model Scenario 0 using Notice → Risk → Act → Explain."),("20 min","Pairs complete Scenarios 1-6; teacher checks reasoning."),("12 min","Independent Scenarios 7-10."),("5 min","Exit ticket and debrief. Honors extension adds 15-25 min.")]
t=doc.add_table(rows=1,cols=2); t.style="Table Grid"
for i,h in enumerate(["Time","Action"]): shade(t.cell(0,i),NAVY); text(t.cell(0,i).paragraphs[0],h,9,True,WHITE)
for a,b in flow:
    c=t.add_row().cells; text(c[0].paragraphs[0],a,9,True,TEAL); text(c[1].paragraphs[0],b,9)
doc.add_paragraph(); box(doc,"QUICK START / SUB PLAN","Assign the Student Tool, all scenario pages, and the exit ticket. Students read the four-step framework independently and complete in order. Review with the key next class; a substitute should not improvise site-specific safety procedures.",GRAY)
box(doc,"DIFFICULTY PATH","Support 1/3: Scenarios 1-6 + category bank (35-45 min)  |  Core 2/3: all 10 (45-55 min)  |  Honors 3/3: all 10 + extension (60-75 min)",GRAY)
box(doc,"COMMON MISCONCEPTIONS","Watch for: reporting only after cleanup; returning excess reagent to stock; treating a colorless liquid as safe; naming a rule without explaining how it reduces the stated risk.",GRAY)
box(doc,"IMPORTANT","This resource supports instruction; it does not replace district policy, teacher supervision, SDS information, or site-specific emergency procedures.","FFF3E0")

# 3 framework
page(doc); title(doc,"Student Tool","Make the Safe Call","Name: __________________________  Class: __________  Date: __________")
banner(doc,"NOTICE → RISK → ACT → EXPLAIN")
steps=[("1  NOTICE","Underline the detail that signals a problem."),("2  RISK","State what could happen or who could be exposed."),("3  ACT","Choose the safest immediate action. Stop first when continuing could increase harm."),("4  EXPLAIN","Connect the action to the evidence. Avoid vague answers such as “be careful.”")]
for h,b in steps: box(doc,h,b)
doc.add_paragraph(); box(doc,"CATEGORY BANK","PPE • chemical handling • heat/flame • glassware • housekeeping • emergency response",GRAY)
p=doc.add_paragraph(); text(p,"Scenario 0 model: ",9,True,TEAL); text(p,"A clearly identified water spill from a wash bottle is visible on a busy walkway. Evidence: the floor is wet. Risk: someone could slip. Action: block traffic and follow the teacher-approved cleanup procedure. Why: the route is protected while the hazard is removed correctly.",9)

# 4-6 scenarios
student_chunks=[(DATA["scenarios"][:4],"Scenarios 1-4"),(DATA["scenarios"][4:7],"Scenarios 5-7"),(DATA["scenarios"][7:],"Scenarios 8-10")]
for chunk, subtitle in student_chunks:
    page(doc); title(doc,"Student Worksheet",subtitle,"Notice the evidence, state the risk, choose the safest first action, and explain why it works.")
    for item in chunk: scenario_card(doc,item)

# 7 honors
page(doc); title(doc,"Honors Extension","Prioritize and defend","Complete after all ten scenarios.")
box(doc,"A. RISK RANKING","Choose the three scenarios that require the fastest intervention. Rank them 1-3. For each, explain severity, likelihood, and how quickly harm could occur.")
for n in range(1,4):
    p=doc.add_paragraph(); text(p,f"Rank {n}: Scenario ____  Reasoning:",9,True,TEAL); add_lines(p._parent,0) if False else None
    for _ in range(2): doc.add_paragraph("________________________________________________________________________________")
box(doc,"B. CLAIM-EVIDENCE-REASONING","Claim: Which single habit would prevent the widest range of incidents in this set? Evidence: cite at least two scenarios. Reasoning: explain why the habit transfers across situations.",GRAY)
for _ in range(8): doc.add_paragraph("________________________________________________________________________________")

# 8 exit ticket
page(doc); title(doc,"Exit Ticket","Ready for the lab?","Name: __________________________  Class: __________")
questions=["1. You notice an unlabeled container at a shared station. What is the safest first action, and why?","2. Why is “be more careful” weaker than a specific safer-action statement?","3. Circle the category that needs more practice: PPE / chemicals / heat / glassware / housekeeping / emergencies","4. Write one question you still have before beginning laboratory work."]
for q in questions:
    box(doc,q,"________________________________________________________________________\n________________________________________________________________________",PALE)
doc.add_paragraph(); box(doc,"SELF-CHECK","I identified evidence [ ]   I stated the risk [ ]   I named a specific action [ ]   I explained why [ ]",GRAY)

# 9-11 keys
key_chunks=[(DATA["scenarios"][:4],"Scenarios 1-4"),(DATA["scenarios"][4:7],"Scenarios 5-7"),(DATA["scenarios"][7:],"Scenarios 8-10 + extension")]
for chunk,subtitle in key_chunks:
    page(doc); title(doc,"Answer Key",subtitle,"Reasonable equivalent wording is acceptable when it follows local policy and does not add risk.")
    for item in chunk: answer_card(doc,item)
    if chunk[0]["id"]==8:
        doc.add_paragraph(); box(doc,"CORE RUBRIC - 3 POINTS PER SCENARIO","1 point: evidence and risk are specific and connected. 1 point: first action is safe, observable, and immediate. 1 point: explanation shows how the action reduces the stated risk.",GRAY)
        box(doc,"HONORS EXEMPLAR","One defensible ranking is 4 first (heated test tube directed at a person), 9 second (chemical remains on skin while response is delayed), and 10 third (energized hot plate near paper). Other rankings earn credit when severity, likelihood, and immediacy are supported with scenario evidence.",GRAY)
        box(doc,"CER MODEL","Claim: Pausing to check the setup before acting prevents the widest range of incidents. Evidence: Scenario 1 begins chemical handling before goggles are in place, and Scenario 6 returns used reagent toward shared stock. Reasoning: a deliberate check catches both personal-protection and contamination risks before exposure spreads.",GRAY)
        box(doc,"EXIT TICKET GUIDANCE","1. Do not touch or test; keep others away and notify the teacher. 2. Specific actions are observable and directly reduce the identified risk. Items 3-4 are formative.")

# 12 sources/terms
page(doc); title(doc,"Teacher Reference","Sources, rights, and use","Version 1.0 • August 2026")
box(doc,"ORIGINALITY","All student scenarios, questions, explanations, and layout in this resource are original. External sources informed factual accuracy and instructional scope; no source questions or passages were copied or adapted.")
refs=[
"American Chemical Society. Guidelines for Chemical Laboratory Safety in Secondary Schools and RAMP resources. https://www.acs.org/education/policies/middle-and-high-school-chemistry/safety.html",
"American Chemical Society. Safety Equipment Guidelines. https://www.acs.org/education/policies/middle-and-high-school-chemistry/classroom-and-lab-facilities/safety-equipment.html",
"American Chemical Society. Student Laboratory Code of Conduct. https://institute.acs.org/acs-center/lab-safety/education-training/high-school-labs/student-lab-code-of-conduct.html",
"OpenStax Chemistry 2e, Chapter 1, used only for scope and terminology benchmarking. Current English edition license is CC BY-NC-SA 4.0; no content was adapted for this commercial resource. https://openstax.org/books/chemistry-2e/pages/preface"
]
for r in refs:
    p=doc.add_paragraph(style=None); p.style=doc.styles["Normal"]; p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.first_line_indent=Inches(-.18); text(p,"• "+r,8)
box(doc,"TERMS OF USE","Single-classroom use by the purchaser. You may print and assign to your own students. Do not resell, post publicly, share the editable file, or upload answer keys to an open website. District and site safety rules take precedence.",GRAY)
box(doc,"QUALITY CHECK","Before student use, confirm that emergency procedures and disposal language match your school, room, and available safety equipment.","FFF3E0")

for s in doc.sections:
    margins(s)
# Section footers remain linked to the first section, so adding the footer to
# every section would duplicate the line once per section in Microsoft Word.
OUT.mkdir(parents=True,exist_ok=True); doc.save(DOCX)
print(DOCX)

# A stable print PDF is built directly so the product does not depend on a local
# Office renderer. The editable DOCX above remains the teacher-editable master.
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas as pdfcanvas

PDF = OUT / "product.pdf"
pdfmetrics.registerFont(TTFont("PoppinsEmbedded", str(ROOT/"fonts/Poppins-Regular.ttf")))
pdfmetrics.registerFont(TTFont("PoppinsEmbedded-Bold", str(ROOT/"fonts/Poppins-Bold.ttf")))
ss = getSampleStyleSheet()
P = ParagraphStyle("Body", parent=ss["BodyText"], fontName="PoppinsEmbedded", fontSize=9, leading=11, textColor=colors.HexColor("#102A43"), spaceAfter=5)
H = ParagraphStyle("Head", parent=ss["Heading1"], fontName="PoppinsEmbedded-Bold", fontSize=20, leading=22, textColor=colors.HexColor("#102A43"), spaceAfter=8)
K = ParagraphStyle("Kick", parent=P, fontName="PoppinsEmbedded-Bold", fontSize=8, textColor=colors.HexColor("#007C83"), spaceAfter=2)
S = ParagraphStyle("Small", parent=P, fontSize=8.5, leading=10.5)
C = ParagraphStyle("Center", parent=P, alignment=TA_CENTER)

def pp(s, style=P): return Paragraph(s.replace("&","&amp;"), style)
def heading(k,h,sub=""):
    out=[pp(k.upper(),K),pp(h,H)]
    if sub: out.append(pp(sub,P))
    return out
def band(s, color="#007C83"):
    t=Table([[pp(s,ParagraphStyle("Band",parent=P,fontName="PoppinsEmbedded-Bold",fontSize=10,textColor=colors.white,alignment=TA_CENTER))]],colWidths=[7.15*inch])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor(color)),("BOX",(0,0),(-1,-1),.5,colors.HexColor(color)),("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6)])); return t
def pbox(head,body,bg="#EAF7F7"):
    t=Table([[pp(head,ParagraphStyle("BH",parent=S,fontName="PoppinsEmbedded-Bold",textColor=colors.white))],[pp(body,S)]],colWidths=[7.15*inch])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(0,0),colors.HexColor("#007C83")),("BACKGROUND",(0,1),(0,1),colors.HexColor(bg)),("BOX",(0,0),(-1,-1),.6,colors.HexColor("#007C83")),("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); return t

story=[]
story += [band("CHEMISTRY FOUNDATIONS • FORMAT 1"),Spacer(1,.45*inch),pp("LAB SAFETY",ParagraphStyle("Cover",parent=H,fontSize=30,leading=32,alignment=TA_CENTER)),pp("SCENARIO ANALYSIS",ParagraphStyle("Cover2",parent=H,fontSize=26,leading=29,textColor=colors.HexColor("#007C83"),alignment=TA_CENTER)),pp("MAKE THE SAFE CALL",ParagraphStyle("Cover3",parent=P,fontName="PoppinsEmbedded-Bold",fontSize=14,textColor=colors.HexColor("#C44536"),alignment=TA_CENTER)),Spacer(1,.3*inch),pbox("WHAT STUDENTS DO","Analyze 10 original laboratory situations, identify evidence of risk, choose the safest first action, and justify decisions using a consistent safety framework."),Spacer(1,.18*inch)]
badges=[[pp(x,ParagraphStyle("Badge",parent=S,fontName="PoppinsEmbedded-Bold",textColor=colors.HexColor("#007C83"),alignment=TA_CENTER)) for x in row] for row in [["GRADES 9-11","CORE 45-55 MIN","PRINT + EDITABLE"],["CORE + HONORS","EXIT TICKET","FULL KEY"]]]
t=Table(badges,colWidths=[2.38*inch]*3,rowHeights=[.42*inch]*2); t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.5,colors.HexColor("#B7C7D3")),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#EAF7F7")),("BACKGROUND",(0,1),(-1,1),colors.HexColor("#E8EEF2")),("VALIGN",(0,0),(-1,-1),"MIDDLE")]))
story += [t,Spacer(1,.5*inch),pp("Original content • Printer-friendly • Structured reasoning",ParagraphStyle("Tag",parent=C,fontName="PoppinsEmbedded-Bold")),PageBreak()]

story += heading("Welcome / Teacher Guide","Welcome, Teacher","Use this ready-to-teach lesson for direct instruction, collaborative practice, a sub plan, or a readiness check.")
story += [pbox("LEARNING TARGET","I can notice evidence, describe the risk, choose the safest immediate action, and explain why it reduces harm."),Spacer(1,8),band("SUGGESTED 50-MINUTE FLOW","#102A43")]
flow_data=[[pp("Time",K),pp("Action",K)]]+[[pp(a,S),pp(b,S)] for a,b in flow]
t=Table(flow_data,colWidths=[.8*inch,6.35*inch]); t.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.4,colors.HexColor("#AABBC8")),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#E8EEF2")),("VALIGN",(0,0),(-1,-1),"TOP"),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5)]))
story += [t,Spacer(1,5),pbox("QUICK START / SUB PLAN","Assign the Student Tool, all scenario pages, and the exit ticket. Students read the four-step framework independently and complete in order. Review with the key next class; a substitute should not improvise site-specific safety procedures.","#E8EEF2"),Spacer(1,4),pbox("DIFFICULTY PATH","Support 1/3: Scenarios 1-6 + category bank (35-45 min)  |  Core 2/3: all 10 (45-55 min)  |  Honors 3/3: all 10 + extension (60-75 min)","#E8EEF2"),Spacer(1,4),pbox("COMMON MISCONCEPTIONS","Reporting only after cleanup; returning excess reagent to stock; treating a colorless liquid as safe; naming a rule without explaining how it reduces the stated risk.","#E8EEF2"),Spacer(1,4),pbox("IMPORTANT","This resource supports instruction; it does not replace district policy, teacher supervision, SDS information, or site-specific emergency procedures.","#FFF3E0"),PageBreak()]

story += heading("Student Tool","Make the Safe Call","Name: __________________________  Class: __________  Date: __________")+[band("NOTICE → RISK → ACT → EXPLAIN")]
for h,b in steps: story += [Spacer(1,5),pbox(h,b)]
story += [Spacer(1,8),pbox("CATEGORY BANK","PPE • chemical handling • heat/flame • glassware • housekeeping • emergency response","#E8EEF2"),Spacer(1,8),pp("<b>Scenario 0 model:</b> A clearly identified water spill from a wash bottle is visible on a busy walkway. Evidence: the floor is wet. Risk: someone could slip. Action: block traffic and follow the teacher-approved cleanup procedure. Why: the route is protected while the hazard is removed correctly."),PageBreak()]

for chunk,subtitle in student_chunks:
    story += heading("Student Worksheet",subtitle,"Notice the evidence, state the risk, choose the safest first action, and explain why it works.")
    for item in chunk:
        body=(f"{xml_escape(item['text'])}<br/>"
              "<b>Evidence noticed:</b> _________________________________________________<br/>"
              "<b>Risk / possible harm:</b> _____________________________________________<br/>"
              "<b>Safest first action:</b> _______________________________________________<br/>"
              "<b>Why this action reduces risk:</b> ______________________________________<br/>"
              "________________________________________________________________________")
        story += [KeepTogether(pbox(f"{item['id']}. {xml_escape(item['title'])}",body)),Spacer(1,5)]
    story.append(PageBreak())

story += heading("Honors Extension","Prioritize and defend","Complete after all ten scenarios.")+[pbox("A. RISK RANKING","Choose the three scenarios that require the fastest intervention. Rank them 1-3. For each, explain severity, likelihood, and how quickly harm could occur.")]
for n in range(1,4): story += [Spacer(1,6),pp(f"<b>Rank {n}: Scenario ____  Reasoning:</b><br/>________________________________________________________________________________<br/>________________________________________________________________________________")]
story += [Spacer(1,8),pbox("B. CLAIM-EVIDENCE-REASONING","Claim: Which single habit would prevent the widest range of incidents in this set? Evidence: cite at least two scenarios. Reasoning: explain why the habit transfers across situations.","#E8EEF2"),pp("<br/>________________________________________________________________________________<br/>________________________________________________________________________________<br/>________________________________________________________________________________<br/>________________________________________________________________________________<br/>________________________________________________________________________________<br/>________________________________________________________________________________"),PageBreak()]

story += heading("Exit Ticket","Ready for the lab?","Name: __________________________  Class: __________")
for q in questions: story += [pbox(q,"________________________________________________________________________<br/>________________________________________________________________________"),Spacer(1,8)]
story += [pbox("SELF-CHECK","I identified evidence [ ]   I stated the risk [ ]   I named a specific action [ ]   I explained why [ ]","#E8EEF2"),PageBreak()]

for items,sub in key_chunks:
    story += heading("Answer Key",sub,"Reasonable equivalent wording is acceptable when it follows local policy and does not add risk.")
    for x in items:
        body=(f"<b>Evidence:</b> {xml_escape(x['evidence'])}<br/><b>Risk:</b> {xml_escape(x['risk'])}<br/>"
              f"<b>Action:</b> {xml_escape(x['action'])}<br/><b>Why:</b> {xml_escape(x['why'])}")
        story += [KeepTogether(pbox(f"{x['id']}. {xml_escape(x['title'])}",body,"#F4F8FA")),Spacer(1,5)]
    if items[0]["id"]==8: story += [Spacer(1,5),pbox("CORE RUBRIC - 3 POINTS PER SCENARIO","1 point: evidence and risk are specific and connected. 1 point: first action is safe, observable, and immediate. 1 point: explanation shows how the action reduces the stated risk.","#E8EEF2"),Spacer(1,4),pbox("HONORS EXEMPLAR","One defensible ranking is 4 first, 9 second, and 10 third. Other rankings earn credit when severity, likelihood, and immediacy are supported with scenario evidence.","#E8EEF2"),Spacer(1,4),pbox("CER MODEL","Claim: Pausing to check the setup before acting prevents the widest range of incidents. Evidence: Scenario 1 begins chemical handling before goggles are in place, and Scenario 6 returns used reagent toward shared stock. Reasoning: a deliberate check catches both personal-protection and contamination risks before exposure spreads.","#E8EEF2"),Spacer(1,4),pbox("EXIT TICKET GUIDANCE","1. Do not touch or test; keep others away and notify the teacher. 2. Specific actions are observable and directly reduce the identified risk. Items 3-4 are formative.")]
    story.append(PageBreak())

story += heading("Teacher Reference","Sources, rights, and use","Version 1.0 • August 2026")+[pbox("ORIGINALITY","All student scenarios, questions, explanations, and layout are original. External sources informed factual accuracy and instructional scope; no source questions or passages were copied or adapted.")]
for r in refs: story += [Spacer(1,5),pp("• "+r,S)]
story += [Spacer(1,8),pbox("TERMS OF USE","Single-classroom use by the purchaser. You may print and assign to your own students. Do not resell, post publicly, share the editable file, or upload answer keys to an open website. District and site safety rules take precedence.","#E8EEF2"),Spacer(1,6),pbox("QUALITY CHECK","Before student use, confirm that emergency procedures and disposal language match your school, room, and available safety equipment.","#FFF3E0")]

def canvas_footer(canvas, doc):
    canvas.saveState(); canvas.setStrokeColor(colors.HexColor("#C7D4DD")); canvas.line(.62*inch,.42*inch,7.88*inch,.42*inch)
    canvas.setFont("PoppinsEmbedded",7); canvas.setFillColor(colors.HexColor("#66788A")); canvas.drawCentredString(4.25*inch,.25*inch,f"Lab Safety Scenario Analysis  |  CurioNest  |  {doc.page}"); canvas.restoreState()

class EmbeddedCanvas(pdfcanvas.Canvas):
    def __init__(self,*args,**kwargs):
        kwargs.setdefault("initialFontName","PoppinsEmbedded")
        kwargs.setdefault("initialFontSize",12)
        super().__init__(*args,**kwargs)

SimpleDocTemplate(str(PDF),pagesize=letter,rightMargin=.62*inch,leftMargin=.62*inch,topMargin=.52*inch,bottomMargin=.55*inch,title=DATA["title"],author="CurioNest").build(story,onFirstPage=canvas_footer,onLaterPages=canvas_footer,canvasmaker=EmbeddedCanvas)
print(PDF)

# Package derivatives and listing images are built from the actual product PDF.
from pypdf import PdfReader, PdfWriter
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import pypdfium2 as pdfium

def subset_pdf(indices, target, watermark=False, footer_label=None):
    reader=PdfReader(str(PDF)); writer=PdfWriter()
    total=len(indices)
    for ordinal,index in enumerate(indices,1):
        page=reader.pages[index]
        if watermark:
            buf=BytesIO(); mark=pdfcanvas.Canvas(buf,pagesize=letter,initialFontName="PoppinsEmbedded")
            mark.saveState(); mark.setFillColor(colors.HexColor("#5A7184")); mark.setFillAlpha(.12)
            mark.setFont("PoppinsEmbedded-Bold",70); mark.translate(4.25*inch,5.5*inch); mark.rotate(35)
            mark.drawCentredString(0,0,"PREVIEW"); mark.restoreState(); mark.save(); buf.seek(0)
            page.merge_page(PdfReader(buf).pages[0])
        if footer_label:
            buf=BytesIO(); foot=pdfcanvas.Canvas(buf,pagesize=letter,initialFontName="PoppinsEmbedded")
            foot.setFillColor(colors.white); foot.rect(0,0,8.5*inch,.48*inch,fill=1,stroke=0)
            foot.setStrokeColor(colors.HexColor("#C7D4DD")); foot.line(.62*inch,.42*inch,7.88*inch,.42*inch)
            foot.setFont("PoppinsEmbedded",7); foot.setFillColor(colors.HexColor("#66788A"))
            foot.drawCentredString(4.25*inch,.25*inch,f"Lab Safety Scenario Analysis  |  {footer_label} {ordinal} of {total}")
            foot.save(); buf.seek(0); page.merge_page(PdfReader(buf).pages[0])
        writer.add_page(page)
    writer.add_metadata({"/Title":DATA["title"],"/Author":"CurioNest"})
    with target.open("wb") as f: writer.write(f)

reader=PdfReader(str(PDF)); page_texts=[" ".join((p.extract_text() or "").upper().split()) for p in reader.pages]
def role(*needles):
    hits=[i for i,t in enumerate(page_texts) if all(n.upper() in t for n in needles)]
    if len(hits)!=1: raise RuntimeError(f"page role {needles} expected once, found {hits}")
    return hits[0]

roles={
    "cover":role("CHEMISTRY FOUNDATIONS","MAKE THE SAFE CALL"),
    "welcome":role("WELCOME, TEACHER","QUICK START / SUB PLAN"),
    "student_tool":role("STUDENT TOOL","MAKE THE SAFE CALL"),
    "scenario_1":role("STUDENT WORKSHEET","SCENARIOS 1-4"),
    "scenario_2":role("STUDENT WORKSHEET","SCENARIOS 5-7"),
    "scenario_3":role("STUDENT WORKSHEET","SCENARIOS 8-10"),
    "honors":role("HONORS EXTENSION","PRIORITIZE AND DEFEND"),
    "exit":role("EXIT TICKET","READY FOR THE LAB"),
    "key_1":role("ANSWER KEY","SCENARIOS 1-4"),
    "key_2":role("ANSWER KEY","SCENARIOS 5-7"),
    "key_3":role("ANSWER KEY","SCENARIOS 8-10 + EXTENSION"),
    "sources":role("TEACHER REFERENCE","SOURCES, RIGHTS, AND USE")
}

STUDENT=OUT/"student-packet.pdf"; PREVIEW=OUT/"preview.pdf"; TEACHER=OUT/"teacher-guide-and-key.pdf"; STUDENT_BW=OUT/"student-packet-bw.pdf"
student_roles=[roles[x] for x in ("student_tool","scenario_1","scenario_2","scenario_3","honors","exit")]
subset_pdf(student_roles,STUDENT,footer_label="Student")
subset_pdf([roles["cover"],roles["scenario_1"],roles["key_1"]],PREVIEW,watermark=True)
subset_pdf([roles[x] for x in ("cover","welcome","key_1","key_2","key_3","sources")],TEACHER,footer_label="Teacher")

student_doc=pdfium.PdfDocument(str(STUDENT)); gray_pages=[]
for pg in student_doc:
    gray_pages.append(pg.render(scale=2.5).to_pil().convert("L").convert("RGB"))
gray_pages[0].save(STUDENT_BW,"PDF",save_all=True,append_images=gray_pages[1:],resolution=180.0,quality=92)

def font(size,bold=False):
    name="Poppins-Bold.ttf" if bold else "Poppins-Regular.ttf"
    return ImageFont.truetype(str(ROOT/"fonts"/name),size)

def contain(img, box, bg="white"):
    copy=img.copy(); copy.thumbnail(box,Image.Resampling.LANCZOS)
    out=Image.new("RGB",box,bg); out.paste(copy,((box[0]-copy.width)//2,(box[1]-copy.height)//2)); return out

# Square cover uses the actual four-step reasoning framework taught inside.
cover=Image.new("RGB",(1800,1800),"#F7FAFC"); d=ImageDraw.Draw(cover)
d.rectangle((0,0,1800,210),fill="#102A43")
d.text((900,105),"CHEMISTRY FOUNDATIONS",font=font(56,True),fill="white",anchor="mm")
d.text((900,365),"LAB SAFETY",font=font(126,True),fill="#102A43",anchor="mm")
d.text((900,505),"SCENARIO ANALYSIS",font=font(96,True),fill="#007C83",anchor="mm")
d.text((900,610),"MAKE THE SAFE CALL",font=font(46,True),fill="#C44536",anchor="mm")
for x,(number,label) in zip([270,690,1110,1530],[("1","NOTICE"),("2","RISK"),("3","ACT"),("4","EXPLAIN")]):
    d.rounded_rectangle((x-165,710,x+165,980),28,fill="#EAF7F7",outline="#007C83",width=5)
    d.ellipse((x-48,745,x+48,841),fill="#007C83"); d.text((x,793),number,font=font(42,True),fill="white",anchor="mm")
    d.text((x,900),label,font=font(37,True),fill="#102A43",anchor="mm")
d.rounded_rectangle((170,1060,1630,1395),30,fill="#EAF7F7",outline="#007C83",width=5)
d.text((900,1150),"10 ORIGINAL LAB SCENARIOS",font=font(54,True),fill="#102A43",anchor="mm")
d.text((900,1250),"Evidence + Risk + Action + Why",font=font(42,True),fill="#007C83",anchor="mm")
d.text((900,1325),"Core + Honors  •  Exit Ticket  •  Full Answer Key",font=font(32,False),fill="#102A43",anchor="mm")
d.rectangle((0,1530,1800,1800),fill="#007C83")
d.text((900,1625),"GRADES 9-11  |  CORE 45-55 MIN",font=font(46,True),fill="white",anchor="mm")
d.text((900,1715),"Print PDF + Editable DOCX",font=font(42,False),fill="white",anchor="mm")
cover.resize((1200,1200),Image.Resampling.LANCZOS).save(OUT/"cover.png",quality=95)

pdfdoc=pdfium.PdfDocument(str(PDF))
page_imgs={name:pdfdoc[index].render(scale=2.0).to_pil().convert("RGB") for name,index in roles.items()}

def crop_content(img, bottom_ratio=.78):
    w,h=img.size
    return img.crop((int(w*.04),int(h*.035),int(w*.96),int(h*bottom_ratio)))

inside=Image.new("RGB",(1800,1800),"#102A43"); di=ImageDraw.Draw(inside)
di.text((900,105),"WHAT'S INSIDE",font=font(74,True),fill="white",anchor="mm")
di.text((900,185),"Teacher-ready directions + authentic student reasoning",font=font(34),fill="#D6F1F1",anchor="mm")
close=contain(crop_content(page_imgs["scenario_1"],.82),(1500,1270),"#102A43")
inside.paste(close,(150,285))
di.rounded_rectangle((250,1585,1550,1735),22,fill="#007C83")
di.text((900,1660),"EVIDENCE • RISK • ACTION • WHY",font=font(42,True),fill="white",anchor="mm")
inside.save(OUT/"listing-02-inside.png",quality=95)

ready=Image.new("RGB",(1800,1800),"#F7FAFC"); dr=ImageDraw.Draw(ready)
dr.text((900,115),"FULL KEY. CLEAR SUPPORT.",font=font(70,True),fill="#102A43",anchor="mm")
dr.text((900,200),"Every scenario keyed with evidence, risk, action, and rationale",font=font(31),fill="#007C83",anchor="mm")
key_close=contain(crop_content(page_imgs["key_1"],.88),(1500,1240),"#F7FAFC")
ready.paste(key_close,(150,300))
dr.rounded_rectangle((190,1600,1610,1740),24,fill="#007C83")
dr.text((900,1670),"LESS PREP • BETTER SAFETY DISCUSSIONS",font=font(43,True),fill="white",anchor="mm")
ready.save(OUT/"listing-03-teacher-ready.png",quality=95)
cover.save(OUT/"listing-01-cover.png",quality=95)

print(STUDENT); print(STUDENT_BW); print(TEACHER); print(PREVIEW); print(OUT/"cover.png")
