"""Build the original Lab Safety Task Cards package from one JSON source."""

from pathlib import Path
import json
import hashlib
from io import BytesIO
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfgen import canvas as pdfcanvas

ROOT = Path(__file__).resolve().parent
PRODUCT_DIR = ROOT / "products" / "lab-safety-task-cards"
SOURCE_DIR = PRODUCT_DIR / "source"
BUYER_OUT = PRODUCT_DIR / "output" / "buyer"
STATIONS_OUT = PRODUCT_DIR / "output" / "stations"
LISTING_OUT = PRODUCT_DIR / "output" / "listing"
DATA = json.loads((SOURCE_DIR / "source.json").read_text(encoding="utf-8"))
CARDS = DATA["cards"]

NAVY = "102A43"
TEAL = "007C83"
GOLD = "D39B1B"
RED = "C44536"
PALE = "EAF7F7"
GRAY = "E8EEF2"
INK = "243B53"
MUTED = "52606D"
WHITE = "FFFFFF"

PORTION_COLORS = {"A": TEAL, "B": "1F6F8B", "C": GOLD, "D": RED}
PORTION_FILL = {"A": PALE, "B": "EAF3F7", "C": "FFF7E0", "D": "FDEEEE"}
LETTER_LABELS = "ABCD"

PORTION_FILES = {
    "A": "portion-a-notice-and-pause.pdf",
    "B": "portion-b-choose-next-move.pdf",
    "C": "portion-c-communicate-procedure.pdf",
    "D": "portion-d-transfer-repair.pdf",
}


def visual_path(card):
    visual = card.get("visual")
    if not visual:
        return None
    path = ROOT / visual["asset"]
    if not path.is_file():
        raise FileNotFoundError(f"Missing verified visual asset for {card['id']}: {path}")
    return path


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Aptos", size=None, color=NAVY, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell_margins(cell)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_border(cell, color="B7C7D3", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_text(paragraph, text, size=9, color=NAVY, bold=False, italic=False):
    return set_run_font(paragraph.add_run(text), size=size, color=color, bold=bold, italic=italic)


def set_document_geometry(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.48)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.header_distance = Inches(0.24)
        section.footer_distance = Inches(0.25)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_paragraph(p)
        add_text(p, "Lab Safety Task Cards  |  CurioNest  |  ", size=7.5, color="66788A")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)


def add_page(doc):
    doc.add_page_break()


def doc_title(doc, kicker, heading, subtitle=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_text(p, kicker.upper(), size=8.5, color=TEAL, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    add_text(p, heading, size=20, color=NAVY, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        add_text(p, subtitle, size=9.5, color=MUTED)


def doc_band(doc, text, fill=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_border(cell, fill, "4")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, text, size=10, color=WHITE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def doc_box(doc, heading, body, fill=PALE, accent=TEAL):
    table = doc.add_table(rows=2, cols=1)
    set_table_geometry(table, [9360])
    top, bottom = table.cell(0, 0), table.cell(1, 0)
    shade(top, accent)
    shade(bottom, fill)
    set_cell_border(top, accent, "4")
    set_cell_border(bottom, accent, "4")
    p = top.paragraphs[0]
    add_text(p, heading, size=8.8, color=WHITE, bold=True)
    p = bottom.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, body, size=8.8, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def card_header_text(card):
    return f"{card['id']}  |  {card['portion']} {card['portion_name'].upper()}  |  {card['level'].upper()}"


def add_card_docx(doc, card, answer=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    accent = PORTION_COLORS[card["portion"]]
    fill = PORTION_FILL[card["portion"]] if not answer else "F4F8FA"
    shade(cell, fill)
    set_cell_border(cell, accent, "10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_text(p, card_header_text(card), size=7.4, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_text(p, card["title"], size=13, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_text(p, card["prompt"], size=9, color=INK)
    cue = visual_path(card) if not answer else None
    if cue:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        picture = p.add_run().add_picture(str(cue), width=Inches(float(card["visual"].get("width_in", 1.1))))
        picture._inline.docPr.set("descr", card["visual"]["alt"])
    for idx, option in enumerate(card["options"]):
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.06)
        p.paragraph_format.space_after = Pt(1)
        add_text(p, f"[{LETTER_LABELS[idx]}] ", size=8.5, color=accent, bold=True)
        add_text(p, option, size=8.5, color=NAVY)
    if not answer:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        add_text(p, "Choice: __________    Confidence:  1  2  3", size=8.5, color=accent, bold=True)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_text(p, "Why? _________________________________________________________________", size=8.5, color=MUTED)
    else:
        fields = [("Correct answer", card["answer"]), ("Evidence", card["evidence"]), ("Risk", card["risk"]), ("Action", card["action"]), ("Why", card["why"]), ("Misconception", card["misconception"])]
        for label, value in fields:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            add_text(p, f"{label}: ", size=7.8, color=accent, bold=True)
            add_text(p, value, size=7.8, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_card_page_docx(doc, cards, answer=False, label="Task Cards"):
    doc_title(doc, label, f"{cards[0]['portion']} - {cards[0]['portion_name']}", "Cut, sort, or use as a two-card partner set. Follow local safety procedures.")
    for card in cards:
        add_card_docx(doc, card, answer=answer)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover_docx(doc):
    doc_band(doc, "CHEMISTRY FOUNDATIONS  |  FORMAT 2", fill=NAVY)
    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "LAB SAFETY", size=30, color=NAVY, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "TASK CARDS", size=30, color=TEAL, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "NOTICE  |  CHOOSE  |  REPORT", size=12, color=RED, bold=True)
    for _ in range(2):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    doc_box(doc, "WHAT STUDENTS DO", "Use 24 original decision cards to notice evidence, choose the safest next move, communicate with the teacher, and repair common safety misconceptions.")
    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    badges = [["GRADES 9-11", "CORE + HONORS", "24 CARDS"], ["PRINT READY", "RECORDING SHEET", "FULL KEY"]]
    for r, row in enumerate(badges):
        for c, value in enumerate(row):
            cell = table.cell(r, c); shade(cell, PALE if r == 0 else GRAY); set_cell_border(cell, "B7C7D3", "6")
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, value, size=8.2, color=TEAL, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(18)
    add_text(p, "Original content  |  Teacher-ready  |  Site-policy aware", size=9.5, color=NAVY, bold=True)


def add_guide_docx(doc):
    add_page(doc); doc_title(doc, "Teacher Guide", "Use the cards flexibly", "One card asks for one decision. The key explains the reasoning and flags the misconception.")
    doc_box(doc, "LEARNING TARGET", "I can notice a safety detail, choose a specific next move, and explain how that move reduces risk.")
    doc_band(doc, "FOUR PORTIONS", fill=NAVY)
    table = doc.add_table(rows=1, cols=3); set_table_geometry(table, [1900, 2600, 4860])
    for i, h in enumerate(["Portion", "Main work", "Use"]):
        cell = table.cell(0, i); shade(cell, GRAY); set_cell_border(cell, "AABBC8", "6"); add_text(cell.paragraphs[0], h, size=8.3, color=NAVY, bold=True)
    rows = [("A", "Notice and pause", "Evidence hunt, labels, PPE, damaged equipment, changed details"), ("B", "Choose the next move", "PPE, food, heat, transport, spill response"), ("C", "Communicate and follow procedure", "Exposure, near miss, unknowns, SDS, defective equipment"), ("D", "Transfer and repair", "Appearance, residual heat, waste, policy, prioritization")]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            shade(cells[i], WHITE); set_cell_border(cells[i], "AABBC8", "6"); add_text(cells[i].paragraphs[0], value, size=8.1, color=NAVY, bold=(i == 0))
    doc.add_paragraph()
    doc_box(doc, "ROTATION OPTIONS", "Pairs: draw one card from each portion and record the choice. Stations: assign one portion per station and rotate every 6-8 minutes. Bell ringer: use one Support card. Honors: require the evidence, risk, action, and why connection.", fill=GRAY, accent=TEAL)
    doc_box(doc, "IMPORTANT", "Teacher/site policy, current SDS, posted emergency instructions, and supervision control. This resource does not give permission to improvise spill cleanup, neutralization, disposal, first aid, or emergency response.", fill="FFF3E0", accent=RED)
    add_page(doc); doc_title(doc, "Teacher Guide", "Fast facilitation and differentiation", "Use the confidence check to decide what to reteach.")
    doc_box(doc, "SUPPORT", "Use the Support cards first. Ask students to underline the evidence, circle the choice, and complete: 'I chose ___ because the detail ___ could cause ___.'.")
    doc_box(doc, "CORE", "Require a choice plus a specific reason. Reject vague answers such as 'be careful' unless the student names the observable action and the stated risk.", fill=GRAY)
    doc_box(doc, "HONORS", "Require comparison, prioritization, or a policy/SDS justification. Accept equivalent wording when it follows the local procedure and does not add risk.", fill="FFF7E0", accent=GOLD)
    doc_band(doc, "COMMON MISCONCEPTIONS TO LISTEN FOR", fill=NAVY)
    for text in [
        "Clear or colorless means harmless.",
        "A spill matters only if someone is visibly injured.",
        "A hot plate is safe the instant the switch is off.",
        "Broken glass can go into ordinary trash.",
        "A generic rule overrides the posted procedure or current SDS.",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3); add_text(p, text, size=9, color=NAVY)


def add_recording_docx(doc):
    for chunk_index in range(2):
        add_page(doc); doc_title(doc, "Student Recording Sheet", f"Cards {chunk_index * 12 + 1}-{chunk_index * 12 + 12}", "Record a choice and a reason. Confidence: 1 = unsure, 2 = developing, 3 = ready.")
        table = doc.add_table(rows=1, cols=4); set_table_geometry(table, [680, 1120, 6040, 1520])
        headers = ["Card", "Choice", "Evidence or reason", "Confidence"]
        for i, h in enumerate(headers):
            cell = table.cell(0, i); shade(cell, NAVY); set_cell_border(cell, NAVY, "6"); p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(p, h, size=7.8, color=WHITE, bold=True)
        for card in CARDS[chunk_index * 12:(chunk_index + 1) * 12]:
            cells = table.add_row().cells
            values = [card["id"], "____", "____________________________________________", " 1  2  3 "]
            for i, value in enumerate(values):
                shade(cells[i], PORTION_FILL[card["portion"]] if i == 0 else WHITE); set_cell_border(cells[i], "AABBC8", "6")
                p = cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 2 else WD_ALIGN_PARAGRAPH.LEFT; add_text(p, value, size=8, color=NAVY, bold=(i == 0))
        doc.add_paragraph(); doc_box(doc, "EXIT REFLECTION", "Which type of decision still needs practice? ____________________________________________________________", fill=GRAY)


def add_sources_docx(doc):
    add_page(doc); doc_title(doc, "Teacher Reference", "Sources, rights, and use", "Version 1.0-draft | August 2026")
    doc_box(doc, "ORIGINALITY", "All student cards, choices, explanations, and layout are original. ACS sources informed factual safety checks only; no passages or figures were copied.")
    doc_box(doc, "SOURCE USE", "Chem Pride / Longwood Flipped Chemistry Classroom was used only as a format benchmark. OpenStax Chemistry 2e was not used as a direct authoring source for this package. Both sources are blocked from commercial adaptation in this product.", fill=GRAY)
    for source in DATA["sources"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        add_text(p, source["name"], size=8.8, color=TEAL, bold=True)
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); add_text(p, source["url"], size=7.8, color=MUTED)
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(7); add_text(p, source["use"], size=8.2, color=NAVY)
    visual_lines = [
        " | ".join(part for part in [source["name"], source.get("creator"), source["license"], source.get("url")] if part)
        for source in DATA.get("visual_sources", [])
    ]
    doc_box(doc, "VISUAL ASSETS", "\n".join(visual_lines), fill="FFF7E0", accent=GOLD)
    doc_box(doc, "TERMS OF USE", "Single-classroom use by the purchaser. Print and assign to your own students. Do not resell, post publicly, share the editable file, or upload the answer key to an open website. District and site safety rules take precedence.", fill=GRAY)


def mark_first_rows_as_headers(doc):
    """Expose the first row of each structured/layout table to screen readers."""
    for table in doc.tables:
        row_properties = table.rows[0]._tr.get_or_add_trPr()
        if row_properties.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            row_properties.append(header)


def build_docx():
    doc = Document()
    set_document_geometry(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].font.color.rgb = rgb(NAVY)
    styles["Normal"].paragraph_format.space_after = Pt(3)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    add_cover_docx(doc)
    add_guide_docx(doc)
    add_recording_docx(doc)
    for start in range(0, len(CARDS), 2):
        add_page(doc); add_card_page_docx(doc, CARDS[start:start + 2], answer=False)
    for start in range(0, len(CARDS), 2):
        add_page(doc); add_card_page_docx(doc, CARDS[start:start + 2], answer=True, label="Answer Key")
    add_sources_docx(doc)
    doc.core_properties.title = DATA["title"]
    doc.core_properties.author = "CurioNest"
    doc.core_properties.subject = "Original laboratory safety decision task cards"
    doc.core_properties.keywords = "lab safety, chemistry, task cards, grades 9-11"
    doc.core_properties.comments = "Original content generated from products/lab-safety-task-cards/source/source.json"
    mark_first_rows_as_headers(doc)
    doc.save(BUYER_OUT / "product-editable.docx")


def setup_pdf_fonts():
    pdfmetrics.registerFont(TTFont("Poppins", str(ROOT / "fonts" / "Poppins-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Poppins-Bold", str(ROOT / "fonts" / "Poppins-Bold.ttf")))


def ptext(text, style):
    return Paragraph(xml_escape(str(text)).replace("\n", "<br/>").replace("  ", "&nbsp;&nbsp;"), style)


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle("Body", parent=base["BodyText"], fontName="Poppins", fontSize=8.6, leading=10.4, textColor=colors.HexColor("#102A43"), spaceAfter=3),
        "small": ParagraphStyle("Small", parent=base["BodyText"], fontName="Poppins", fontSize=7.5, leading=9, textColor=colors.HexColor("#243B53"), spaceAfter=2),
        "tiny": ParagraphStyle("Tiny", parent=base["BodyText"], fontName="Poppins", fontSize=6.8, leading=8, textColor=colors.HexColor("#243B53"), spaceAfter=1),
        "head": ParagraphStyle("Head", parent=base["Heading1"], fontName="Poppins-Bold", fontSize=20, leading=23, textColor=colors.HexColor("#102A43"), spaceAfter=5),
        "sub": ParagraphStyle("Sub", parent=base["BodyText"], fontName="Poppins", fontSize=9.5, leading=11.5, textColor=colors.HexColor("#52606D"), spaceAfter=6),
        "kicker": ParagraphStyle("Kicker", parent=base["BodyText"], fontName="Poppins-Bold", fontSize=8, leading=9, textColor=colors.HexColor("#007C83"), spaceAfter=2),
        "band": ParagraphStyle("Band", parent=base["BodyText"], fontName="Poppins-Bold", fontSize=9, leading=11, textColor=colors.white, alignment=TA_CENTER),
        "card_header": ParagraphStyle("CardHeader", parent=base["BodyText"], fontName="Poppins-Bold", fontSize=6.8, leading=8, textColor=colors.HexColor("#007C83"), spaceAfter=2),
        "card_title": ParagraphStyle("CardTitle", parent=base["BodyText"], fontName="Poppins-Bold", fontSize=11.3, leading=13, textColor=colors.HexColor("#102A43"), spaceAfter=3),
        "card_body": ParagraphStyle("CardBody", parent=base["BodyText"], fontName="Poppins", fontSize=8.1, leading=9.8, textColor=colors.HexColor("#243B53"), spaceAfter=2),
        "card_option": ParagraphStyle("CardOption", parent=base["BodyText"], fontName="Poppins", fontSize=7.5, leading=8.7, textColor=colors.HexColor("#102A43"), leftIndent=8, firstLineIndent=-8, spaceAfter=1),
        "key": ParagraphStyle("Key", parent=base["BodyText"], fontName="Poppins", fontSize=7.1, leading=8.4, textColor=colors.HexColor("#243B53"), spaceAfter=1),
    }


def pdf_band(text, style, fill=NAVY):
    return Table([[Paragraph(xml_escape(text), style)]], colWidths=[7.6 * inch], style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + fill)), ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#" + fill)), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))


def pdf_box(head, body, styles, fill=PALE, accent=TEAL):
    h = Paragraph(xml_escape(head), ParagraphStyle("BoxHead", parent=styles["small"], fontName="Poppins-Bold", textColor=colors.white))
    b = Paragraph(xml_escape(body), styles["small"])
    return Table([[h], [b]], colWidths=[7.6 * inch], style=TableStyle([("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#" + accent)), ("BACKGROUND", (0, 1), (0, 1), colors.HexColor("#" + fill)), ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#" + accent)), ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))


def pdf_title(story, styles, kicker, title, subtitle=""):
    story.append(Paragraph(xml_escape(kicker.upper()), styles["kicker"]))
    story.append(Paragraph(xml_escape(title), styles["head"]))
    if subtitle:
        story.append(Paragraph(xml_escape(subtitle), styles["sub"]))


def pdf_cover(story, styles):
    story.append(pdf_band("CHEMISTRY FOUNDATIONS  |  FORMAT 2", styles["band"], NAVY)); story.append(Spacer(1, 0.38 * inch))
    story.append(Paragraph("LAB SAFETY", ParagraphStyle("CoverA", parent=styles["head"], fontSize=31, leading=33, alignment=TA_CENTER, textColor=colors.HexColor("#102A43"))))
    story.append(Paragraph("TASK CARDS", ParagraphStyle("CoverB", parent=styles["head"], fontSize=29, leading=31, alignment=TA_CENTER, textColor=colors.HexColor("#007C83"))))
    story.append(Paragraph("NOTICE  |  CHOOSE  |  REPORT", ParagraphStyle("CoverC", parent=styles["sub"], fontName="Poppins-Bold", fontSize=12, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#C44536"))))
    story.append(Spacer(1, 0.25 * inch))
    story.append(pdf_box("WHAT STUDENTS DO", "Use 24 original decision cards to notice evidence, choose the safest next move, communicate with the teacher, and repair common safety misconceptions.", styles))
    story.append(Spacer(1, 0.18 * inch))
    badge_rows = [[ptext(x, styles["small"]) for x in row] for row in [["GRADES 9-11", "CORE + HONORS", "24 CARDS"], ["PRINT READY", "RECORDING SHEET", "FULL KEY"]]]
    badge = Table(badge_rows, colWidths=[2.53 * inch] * 3, rowHeights=[0.38 * inch] * 2, style=TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B7C7D3")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF7F7")), ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#E8EEF2")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (-1, -1), "CENTER")]))
    story.append(badge); story.append(Spacer(1, 0.38 * inch)); story.append(Paragraph("Original content  |  Teacher-ready  |  Site-policy aware", ParagraphStyle("CoverTag", parent=styles["body"], fontName="Poppins-Bold", alignment=TA_CENTER)))
    story.append(PageBreak())


def pdf_guide(story, styles):
    pdf_title(story, styles, "Teacher Guide", "Use the cards flexibly", "One card asks for one decision. The key explains the reasoning and flags the misconception.")
    story.append(pdf_box("LEARNING TARGET", "I can notice a safety detail, choose a specific next move, and explain how that move reduces risk.", styles)); story.append(Spacer(1, 7))
    story.append(pdf_band("FOUR PORTIONS", styles["band"], NAVY)); story.append(Spacer(1, 5))
    rows = [[ptext(x, styles["small"]) for x in row] for row in [["Portion", "Main work", "Use"], ["A", "Notice and pause", "Evidence, labels, PPE, damaged equipment, changed details"], ["B", "Choose the next move", "PPE, food, heat, transport, spill response"], ["C", "Communicate and follow procedure", "Exposure, near miss, unknowns, SDS, defective equipment"], ["D", "Transfer and repair", "Appearance, residual heat, waste, policy, prioritization"]]]
    t = Table(rows, colWidths=[0.65 * inch, 1.85 * inch, 5.1 * inch], style=TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AABBC8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF2")), ("FONTNAME", (0, 0), (-1, 0), "Poppins-Bold"), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)])); story.append(t); story.append(Spacer(1, 8))
    story.append(pdf_box("ROTATION OPTIONS", "Pairs: draw one card from each portion and record the choice. Stations: assign one portion per station and rotate every 6-8 minutes. Bell ringer: use one Support card. Honors: require the evidence, risk, action, and why connection.", styles, fill=GRAY)); story.append(Spacer(1, 5))
    story.append(pdf_box("IMPORTANT", "Teacher/site policy, current SDS, posted emergency instructions, and supervision control. This resource does not give permission to improvise spill cleanup, neutralization, disposal, first aid, or emergency response.", styles, fill="FFF3E0", accent=RED)); story.append(PageBreak())
    pdf_title(story, styles, "Teacher Guide", "Fast facilitation and differentiation", "Use the confidence check to decide what to reteach.")
    story.append(pdf_box("SUPPORT", "Use the Support cards first. Ask students to underline the evidence, circle the choice, and complete: 'I chose ___ because the detail ___ could cause ___.'.", styles)); story.append(Spacer(1, 5))
    story.append(pdf_box("CORE", "Require a choice plus a specific reason. Reject vague answers such as 'be careful' unless the student names the observable action and the stated risk.", styles, fill=GRAY)); story.append(Spacer(1, 5))
    story.append(pdf_box("HONORS", "Require comparison, prioritization, or a policy/SDS justification. Accept equivalent wording when it follows the local procedure and does not add risk.", styles, fill="FFF7E0", accent=GOLD)); story.append(Spacer(1, 7))
    story.append(pdf_band("COMMON MISCONCEPTIONS TO LISTEN FOR", styles["band"], NAVY)); story.append(Spacer(1, 5))
    for item in ["Clear or colorless means harmless.", "A spill matters only if someone is visibly injured.", "A hot plate is safe the instant the switch is off.", "Broken glass can go into ordinary trash.", "A generic rule overrides the posted procedure or current SDS."]:
        story.append(Paragraph("- " + xml_escape(item), styles["body"]))
    story.append(PageBreak())


def pdf_recording(story, styles):
    for chunk_index in range(2):
        pdf_title(story, styles, "Student Recording Sheet", f"Cards {chunk_index * 12 + 1}-{chunk_index * 12 + 12}", "Record a choice and a reason. Confidence: 1 = unsure, 2 = developing, 3 = ready.")
        rows = [[ptext(x, styles["tiny"]) for x in ["Card", "Choice", "Evidence or reason", "Confidence"]]]
        for card in CARDS[chunk_index * 12:(chunk_index + 1) * 12]:
            rows.append([ptext(card["id"], styles["tiny"]), ptext("____", styles["tiny"]), ptext("________________________________________", styles["tiny"]), ptext("1  2  3", styles["tiny"])])
        t = Table(rows, colWidths=[0.55 * inch, 0.85 * inch, 4.9 * inch, 1.3 * inch], rowHeights=[0.3 * inch] + [0.43 * inch] * 12, style=TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AABBC8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#102A43")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (1, -1), "CENTER"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        for row_index, card in enumerate(CARDS[chunk_index * 12:(chunk_index + 1) * 12], start=1):
            t.setStyle(TableStyle([("BACKGROUND", (0, row_index), (0, row_index), colors.HexColor("#" + PORTION_FILL[card["portion"]]))]))
        story.append(t); story.append(Spacer(1, 10)); story.append(pdf_box("EXIT REFLECTION", "Which type of decision still needs practice? ____________________________________________________________", styles, fill=GRAY)); story.append(PageBreak())


def pdf_card(story, styles, card, answer=False):
    accent = PORTION_COLORS[card["portion"]]
    fill = "F4F8FA" if answer else PORTION_FILL[card["portion"]]
    flows = [Paragraph(xml_escape(card_header_text(card)), styles["card_header"]), Paragraph(xml_escape(card["title"]), styles["card_title"]), Paragraph(xml_escape(card["prompt"]), styles["card_body"])]
    cue = visual_path(card) if not answer else None
    if cue:
        from PIL import Image as PILImage
        with PILImage.open(cue) as source_image:
            pixel_width, pixel_height = source_image.size
        display_width = float(card["visual"].get("width_in", 1.1)) * inch
        display_height = display_width * pixel_height / pixel_width
        visual = RLImage(str(cue), width=display_width, height=display_height)
        visual.hAlign = "CENTER"
        flows.extend([visual, Spacer(1, 2)])
    for idx, option in enumerate(card["options"]):
        flows.append(Paragraph(f"<b>{LETTER_LABELS[idx]}.</b> {xml_escape(option)}", styles["card_option"]))
    if not answer:
        flows.extend([Spacer(1, 4), Paragraph("<b>Choice:</b> ____________________    <b>Confidence:</b> 1  2  3", styles["card_body"]), Paragraph("<b>Why?</b> __________________________________________________________", styles["card_body"])])
    else:
        for label, value in [("Correct answer", card["answer"]), ("Evidence", card["evidence"]), ("Risk", card["risk"]), ("Action", card["action"]), ("Why", card["why"]), ("Misconception", card["misconception"])]:
            flows.append(Paragraph(f"<font color='#{accent}'><b>{xml_escape(label)}:</b></font> {xml_escape(value)}", styles["key"]))
    card_table = Table([[flows]], colWidths=[7.6 * inch], rowHeights=[4.08 * inch], style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + fill)), ("BOX", (0, 0), (-1, -1), 1.0, colors.HexColor("#" + accent)), ("LEFTPADDING", (0, 0), (-1, -1), 9), ("RIGHTPADDING", (0, 0), (-1, -1), 9), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 6), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    return card_table


def pdf_cards(story, styles, answer=False, cards=None):
    selected = CARDS if cards is None else cards
    for start in range(0, len(selected), 2):
        group = selected[start:start + 2]
        pdf_title(story, styles, "Answer Key" if answer else "Student Task Cards", f"{group[0]['portion']} - {group[0]['portion_name']}", "Reason from evidence. Follow local safety procedures.")
        story.append(pdf_card(story, styles, group[0], answer=answer)); story.append(Spacer(1, 0.12 * inch)); story.append(pdf_card(story, styles, group[1], answer=answer)); story.append(PageBreak())


def pdf_sources(story, styles):
    pdf_title(story, styles, "Teacher Reference", "Sources, rights, and use", "Version 1.0-draft | August 2026")
    story.append(pdf_box("ORIGINALITY", "All student cards, choices, explanations, and layout are original. ACS sources informed factual safety checks only; no passages or figures were copied.", styles)); story.append(Spacer(1, 6))
    story.append(pdf_box("SOURCE USE", "Chem Pride / Longwood Flipped Chemistry Classroom was used only as a format benchmark. OpenStax Chemistry 2e was not used as a direct authoring source for this package. Both sources are blocked from commercial adaptation in this product.", styles, fill=GRAY)); story.append(Spacer(1, 6))
    for source in DATA["sources"]:
        story.append(Paragraph(xml_escape(source["name"]), styles["kicker"])); story.append(Paragraph(xml_escape(source["url"]), styles["small"])); story.append(Paragraph(xml_escape(source["use"]), styles["small"])); story.append(Spacer(1, 5))
    visual_text = "; ".join(
        " | ".join(part for part in [source["name"], source.get("creator"), source["license"], source.get("url")] if part)
        for source in DATA.get("visual_sources", [])
    )
    story.append(pdf_box("VISUAL ASSETS", visual_text, styles, fill="FFF7E0", accent=GOLD)); story.append(Spacer(1, 6))
    story.append(pdf_box("TERMS OF USE", "Single-classroom use by the purchaser. Print and assign to your own students. Do not resell, post publicly, share the editable file, or upload the answer key to an open website. District and site safety rules take precedence.", styles, fill=GRAY)); story.append(PageBreak())


def pdf_portion_cover(story, styles, cards):
    portion = cards[0]["portion"]
    accent = PORTION_COLORS[portion]
    story.append(pdf_band(f"LAB SAFETY TASK CARDS  |  PORTION {portion}", styles["band"], accent))
    story.append(Spacer(1, 0.22 * inch))
    story.append(Paragraph(xml_escape(cards[0]["portion_name"]), ParagraphStyle("PortionCover", parent=styles["head"], fontSize=27, leading=31, alignment=TA_CENTER, textColor=colors.HexColor("#" + accent))))
    story.append(Paragraph("6-card station pack  |  Student cards + recording + key", ParagraphStyle("PortionSub", parent=styles["sub"], alignment=TA_CENTER)))
    story.append(Spacer(1, 0.12 * inch))
    focus = {
        "A": "Notice the observable detail that should make the group pause.",
        "B": "Choose the safest immediate move from realistic alternatives.",
        "C": "Report clearly and use labels, SDS, posted procedure, and teacher direction.",
        "D": "Repair unsafe reasoning and transfer the decision to a changed case.",
    }[portion]
    story.append(pdf_box("STATION FOCUS", focus, styles, fill=PORTION_FILL[portion], accent=accent))
    story.append(Spacer(1, 0.15 * inch))
    rows = [[ptext(x, styles["tiny"]) for x in ["Card", "Choice", "Evidence or reason", "Confidence"]]]
    for card in cards:
        rows.append([ptext(card["id"], styles["tiny"]), ptext("____", styles["tiny"]), ptext("________________________________________", styles["tiny"]), ptext("1  2  3", styles["tiny"])])
    table = Table(rows, colWidths=[0.55 * inch, 0.85 * inch, 4.9 * inch, 1.3 * inch], rowHeights=[0.3 * inch] + [0.48 * inch] * 6, style=TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#AABBC8")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + accent)), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (1, -1), "CENTER"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    for row_index in range(1, len(rows)):
        table.setStyle(TableStyle([("BACKGROUND", (0, row_index), (0, row_index), colors.HexColor("#" + PORTION_FILL[portion]))]))
    story.append(table)
    story.append(Spacer(1, 8))
    story.append(pdf_box("USE", "Print pages 2-4 for students. Keep pages 5-7 as the teacher key. Page 8 records sources, rights, and use.", styles, fill=GRAY, accent=NAVY))
    story.append(PageBreak())


def page_footer(canvas, doc):
    canvas.saveState(); canvas.setStrokeColor(colors.HexColor("#C7D4DD")); canvas.line(0.45 * inch, 0.36 * inch, 8.05 * inch, 0.36 * inch); canvas.setFont("Poppins", 7); canvas.setFillColor(colors.HexColor("#66788A")); canvas.drawCentredString(4.25 * inch, 0.2 * inch, f"Lab Safety Task Cards  |  CurioNest  |  {doc.page}"); canvas.restoreState()


def build_pdf(path, include_guide=False, include_recording=False, include_cards=False, include_key=False, include_sources=False):
    setup_pdf_fonts(); styles = pdf_styles(); story = []
    pdf_cover(story, styles)
    if include_guide: pdf_guide(story, styles)
    if include_recording: pdf_recording(story, styles)
    if include_cards: pdf_cards(story, styles, answer=False)
    if include_key: pdf_cards(story, styles, answer=True)
    if include_sources: pdf_sources(story, styles)
    SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.42 * inch, bottomMargin=0.48 * inch, title=DATA["title"], author="CurioNest").build(story, onFirstPage=page_footer, onLaterPages=page_footer)


def build_portion_pdf(path, cards):
    setup_pdf_fonts(); styles = pdf_styles(); story = []
    pdf_portion_cover(story, styles, cards)
    pdf_cards(story, styles, answer=False, cards=cards)
    pdf_cards(story, styles, answer=True, cards=cards)
    pdf_sources(story, styles)
    title = f"Lab Safety Task Cards - Portion {cards[0]['portion']}: {cards[0]['portion_name']}"
    SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.42 * inch, bottomMargin=0.48 * inch, title=title, author="CurioNest").build(story, onFirstPage=page_footer, onLaterPages=page_footer)


def make_preview():
    from pypdf import PdfReader, PdfWriter
    product = PdfReader(str(BUYER_OUT / "product.pdf")); writer = PdfWriter()
    for index in [0, 5, 18]:
        if index < len(product.pages): writer.add_page(product.pages[index])
    with (LISTING_OUT / "preview.pdf").open("wb") as stream: writer.write(stream)


def make_bw_student():
    import pypdfium2 as pdfium
    from PIL import Image
    source = pdfium.PdfDocument(str(BUYER_OUT / "student-packet.pdf")); images = [page.render(scale=2.0).to_pil().convert("L").convert("RGB") for page in source]
    images[0].save(BUYER_OUT / "student-packet-bw.pdf", "PDF", save_all=True, append_images=images[1:], resolution=144.0)


def make_listing_assets():
    import pypdfium2 as pdfium
    from PIL import Image, ImageDraw, ImageFont
    doc = pdfium.PdfDocument(str(BUYER_OUT / "product.pdf"))
    rendered = [doc[i].render(scale=1.8).to_pil().convert("RGB") for i in range(min(8, len(doc)))]
    def font(size, bold=False): return ImageFont.truetype(str(ROOT / "fonts" / ("Poppins-Bold.ttf" if bold else "Poppins-Regular.ttf")), size)
    cover = Image.new("RGB", (1800, 1800), "#F7FAFC"); d = ImageDraw.Draw(cover)
    d.rectangle((0, 0, 1800, 220), fill="#102A43"); d.text((900, 110), "CHEMISTRY FOUNDATIONS", font=font(56, True), fill="white", anchor="mm")
    d.text((900, 390), "LAB SAFETY", font=font(128, True), fill="#102A43", anchor="mm"); d.text((900, 535), "TASK CARDS", font=font(108, True), fill="#007C83", anchor="mm"); d.text((900, 655), "NOTICE  |  CHOOSE  |  REPORT", font=font(42, True), fill="#C44536", anchor="mm")
    for x, label in zip([310, 760, 1210, 1660], ["NOTICE", "CHOOSE", "REPORT", "REPAIR"]):
        d.rounded_rectangle((x - 150, 800, x + 150, 1040), 25, fill="#EAF7F7", outline="#007C83", width=5); d.text((x, 920), label, font=font(31, True), fill="#102A43", anchor="mm")
    d.rounded_rectangle((180, 1150, 1620, 1430), 28, fill="#EAF7F7", outline="#007C83", width=5); d.text((900, 1245), "24 ORIGINAL DECISION CARDS", font=font(54, True), fill="#102A43", anchor="mm"); d.text((900, 1350), "Support  |  Core  |  Honors  |  Full Key", font=font(36), fill="#007C83", anchor="mm")
    d.rectangle((0, 1570, 1800, 1800), fill="#007C83"); d.text((900, 1645), "GRADES 9-11  |  PRINT + EDITABLE", font=font(43, True), fill="white", anchor="mm"); d.text((900, 1730), "Recording sheet  |  Teacher guide  |  Preview", font=font(31), fill="white", anchor="mm")
    cover.resize((1200, 1200), Image.Resampling.LANCZOS).save(LISTING_OUT / "cover.png", quality=95)
    cover.save(LISTING_OUT / "listing-01-cover.png", quality=95)
    def contain(image, box):
        copy = image.copy(); copy.thumbnail(box, Image.Resampling.LANCZOS); out = Image.new("RGB", box, "#102A43"); out.paste(copy, ((box[0] - copy.width) // 2, (box[1] - copy.height) // 2)); return out
    inside = Image.new("RGB", (1800, 1800), "#102A43"); di = ImageDraw.Draw(inside); di.text((900, 105), "WHAT'S INSIDE", font=font(70, True), fill="white", anchor="mm"); di.text((900, 190), "Short decisions with reasons students can defend", font=font(31), fill="#D6F1F1", anchor="mm"); inside.paste(contain(rendered[5], (1500, 1250)), (150, 300)); di.rounded_rectangle((240, 1600, 1560, 1740), 24, fill="#007C83"); di.text((900, 1670), "NOTICE  |  CHOOSE  |  REPORT", font=font(42, True), fill="white", anchor="mm"); inside.save(LISTING_OUT / "listing-02-inside.png", quality=95)
    ready = Image.new("RGB", (1800, 1800), "#F7FAFC"); dr = ImageDraw.Draw(ready); dr.text((900, 115), "FULL KEY. CLEAR SUPPORT.", font=font(65, True), fill="#102A43", anchor="mm"); dr.text((900, 200), "Evidence, risk, action, why, and misconception notes", font=font(29), fill="#007C83", anchor="mm"); ready.paste(contain(rendered[-1], (1500, 1230)), (150, 300)); dr.rounded_rectangle((190, 1600, 1610, 1740), 24, fill="#007C83"); dr.text((900, 1670), "TEACHER-READY ANSWER KEY", font=font(43, True), fill="white", anchor="mm"); ready.save(LISTING_OUT / "listing-03-teacher-ready.png", quality=95)


def main():
    for directory in [SOURCE_DIR, BUYER_OUT, STATIONS_OUT, LISTING_OUT]:
        directory.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256(json.dumps(CARDS, sort_keys=True).encode("utf-8")).hexdigest()
    DATA["review"]["cards_sha256"] = digest
    (SOURCE_DIR / "source.json").write_text(json.dumps(DATA, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    build_docx()
    build_pdf(BUYER_OUT / "product.pdf", include_guide=True, include_recording=True, include_cards=True, include_key=True, include_sources=True)
    build_pdf(BUYER_OUT / "student-packet.pdf", include_recording=True, include_cards=True)
    build_pdf(BUYER_OUT / "teacher-guide-and-key.pdf", include_guide=True, include_key=True, include_sources=True)
    for portion, filename in PORTION_FILES.items():
        build_portion_pdf(STATIONS_OUT / filename, [card for card in CARDS if card["portion"] == portion])
    make_preview(); make_bw_student(); make_listing_assets()
    print("Built:", PRODUCT_DIR / "output")
    for name in ["product.pdf", "product-editable.docx", "student-packet.pdf", "student-packet-bw.pdf", "teacher-guide-and-key.pdf"]:
        print(BUYER_OUT / name)
    for name in PORTION_FILES.values():
        print(STATIONS_OUT / name)
    for name in ["preview.pdf", "cover.png", "listing-01-cover.png", "listing-02-inside.png", "listing-03-teacher-ready.png"]:
        print(LISTING_OUT / name)


if __name__ == "__main__":
    main()
