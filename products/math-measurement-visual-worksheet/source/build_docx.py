"""Build an editable US Letter DOCX companion for the TPT resource."""

from pathlib import Path
import json

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
PRODUCT = HERE.parent
ASSETS = PRODUCT / "assets"
OUT = PRODUCT / "output" / "buyer-files"
DATA = json.loads((HERE / "source.json").read_text(encoding="utf-8"))
BRAND = DATA["brand"]
COPYRIGHT = DATA["source_policy"]["buyer_facing_copyright"]

NAVY = "203B4D"
ORANGE = "C86543"
WASH = "F8F5EE"
NOTE = "F3E5BD"
MUTED = "6D7478"
LINE = "D8D5CE"
WHITE = "FFFEFA"
FONT = "Aptos"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_run(run, size=9.5, bold=False, color="202B33", italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(container, text, size=9.5, bold=False, color="202B33", before=0, after=3, align=None):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    return p


def add_cell_text(cell, text, size=8.5, bold=False, color="202B33", align=None):
    first = cell.paragraphs[0]
    has_drawing = bool(first._p.xpath(".//w:drawing"))
    p = first if len(cell.paragraphs) == 1 and not first.text and not has_drawing else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_answer(container, text):
    table = container.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [7.18])
    set_cell_shading(table.cell(0, 0), NOTE)
    clear_cell(table.cell(0, 0))
    add_cell_text(table.cell(0, 0), text, 8.5, True, ORANGE)
    return table


def add_lines(container, count=2):
    for _ in range(count):
        add_text(container, "________________________________________________________________________________", 8, color=MUTED, after=1)


def add_header(doc, page, title, subtitle, key=False):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [0.62, 6.56])
    set_cell_shading(table.cell(0, 0), ORANGE)
    p = clear_cell(table.cell(0, 0))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(str(page)), 15, True, WHITE)
    clear_cell(table.cell(0, 1))
    add_cell_text(table.cell(0, 1), title, 18, True, NAVY)
    add_cell_text(table.cell(0, 1), subtitle, 7.5, False, MUTED)
    if key:
        add_cell_text(table.cell(0, 1), "TEACHER KEY", 7.5, True, ORANGE, WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        add_cell_text(
            table.cell(0, 1),
            "Name: ____________________   Date: __________   Class Period: ______",
            7.2,
            False,
            MUTED,
            WD_ALIGN_PARAGRAPH.RIGHT,
        )
    add_text(doc, "", 2, after=1)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{COPYRIGHT}  |  Editable Edition"), 7, color=MUTED)


def page1(doc, key=False):
    add_header(doc, 1, "Measurement Ready", "Use the tool, scale, unit, and calculation as one evidence chain.", key)
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [7.18])
    set_cell_shading(table.cell(0, 0), WASH)
    clear_cell(table.cell(0, 0))
    add_cell_text(table.cell(0, 0), "LEARNING TARGETS", 8.5, True, ORANGE)
    add_cell_text(table.cell(0, 0), " • ".join(DATA["learning_targets"][:3]), 8.2)
    add_text(doc, "Three tools, three different jobs", 11.5, True, NAVY, before=4, after=3)
    tools = doc.add_table(rows=1, cols=3)
    tools.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(tools, [2.39, 2.39, 2.40])
    entries = [
        ("electronic-balance-ccby3.png", "Electronic balance", "Measures mass", "Typical unit: g"),
        ("graduated-cylinder-ccby3.png", "Graduated cylinder", "Measures liquid volume", "Typical unit: mL"),
        ("beaker-ccby3.png", "Beaker", "Holds and mixes", "Volume is approximate"),
    ]
    for cell, (filename, name, job, unit) in zip(tools.rows[0].cells, entries):
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ASSETS / filename), height=Inches(1.0))
        add_cell_text(cell, name, 8.7, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(cell, job, 7.6, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(cell, unit, 7.6, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "Measurement rules that prevent avoidable errors", 11.5, True, NAVY, before=4, after=2)
    for item in [
        "Record the value and the unit together.",
        "For an analog scale, record all certain digits plus one estimated digit.",
        "Read a concave liquid meniscus at eye level from its lowest point.",
        "Keep units through every calculation and round only the final result.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        set_run(p.add_run(item), 8.7)
    formula = doc.add_table(rows=1, cols=2)
    set_table_widths(formula, [4.65, 2.53])
    for cell in formula.rows[0].cells:
        set_cell_shading(cell, NOTE)
        clear_cell(cell)
    add_cell_text(formula.cell(0, 0), "FORMULA TOOLBOX", 8.3, True, ORANGE)
    add_cell_text(formula.cell(0, 0), "density = mass / volume\npercent error =\n|experimental - accepted| / accepted × 100", 8.8, True, NAVY)
    add_cell_text(formula.cell(0, 1), "UNITS", 8.3, True, ORANGE)
    add_cell_text(formula.cell(0, 1), "mass: g\nliquid volume: mL\ndensity: g/mL", 8.5)


def page2(doc, key=False):
    add_header(doc, 2, "Choose the Tool", "Name what you see, then use visible features to justify the choice.", key)
    add_text(doc, "A. Identify each instrument", 11.5, True, NAVY, after=3)
    tools = doc.add_table(rows=1, cols=3)
    set_table_widths(tools, [2.39, 2.39, 2.40])
    entries = [
        ("electronic-balance-ccby3.png", "1", DATA["answers"]["tool_1"]),
        ("graduated-cylinder-ccby3.png", "2", DATA["answers"]["tool_2"]),
        ("beaker-ccby3.png", "3", DATA["answers"]["tool_3"]),
    ]
    for cell, (filename, number, answer) in zip(tools.rows[0].cells, entries):
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(number + "\n"), 8.5, True, ORANGE)
        p.add_run().add_picture(str(ASSETS / filename), height=Inches(0.92))
        if key:
            set_cell_shading(cell, NOTE)
            add_cell_text(cell, answer.replace(" - ", "\n"), 7.8, True, ORANGE)
        else:
            add_cell_text(cell, "Instrument: __________________", 7.6)
            add_cell_text(cell, "Quantity/job: ________________", 7.6)
            add_cell_text(cell, "Unit: _______________________", 7.6)
    add_text(doc, "B. Select and justify", 11.5, True, NAVY, before=4, after=2)
    add_text(doc, "4. You need 18.6 mL of water. Choose a beaker or a graduated cylinder. Explain which visible design features support your choice.", 8.8)
    if key:
        add_answer(doc, DATA["answers"]["choice_1"])
    else:
        add_lines(doc, 2)
    add_text(doc, "5. You need about 125 mL of solution and room to stir. Choose a beaker or a graduated cylinder. Explain.", 8.8, before=3)
    if key:
        add_answer(doc, DATA["answers"]["choice_2"])
    else:
        add_lines(doc, 2)
    add_text(doc, "C. Read visual evidence", 11.5, True, NAVY, before=4, after=2)
    evidence = doc.add_table(rows=1, cols=2)
    set_table_widths(evidence, [1.55, 5.63])
    clear_cell(evidence.cell(0, 0))
    p = evidence.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / "graduated-cylinder-public-domain.jpg"), height=Inches(1.55))
    clear_cell(evidence.cell(0, 1))
    add_cell_text(evidence.cell(0, 1), "6. Identify two visible features that show this vessel was designed to measure liquid volume.", 8.8)
    if key:
        set_cell_shading(evidence.cell(0, 1), NOTE)
        add_cell_text(evidence.cell(0, 1), DATA["answers"]["photo_evidence"], 8.3, True, ORANGE)
    else:
        add_cell_text(evidence.cell(0, 1), "\n____________________________________________________________\n\n____________________________________________________________", 8, False, MUTED)
    add_text(doc, "Photo: Darrien / Haltopub, Public Domain, Wikimedia Commons", 6.5, color=MUTED, after=0)


def page3(doc, key=False):
    add_header(doc, 3, "Calculate with Units", "Show the formula, substitute values with units, and round only the final answer.", key)
    guided = doc.add_table(rows=1, cols=1)
    set_table_widths(guided, [7.18])
    set_cell_shading(guided.cell(0, 0), WASH)
    clear_cell(guided.cell(0, 0))
    add_cell_text(guided.cell(0, 0), "GUIDED EXAMPLE", 8.3, True, ORANGE)
    add_cell_text(guided.cell(0, 0), "A sample has a mass of 36.4 g and a volume of 14.0 mL.\nd = m / V = 36.4 g / 14.0 mL = 2.60 g/mL", 9.2)
    tasks = [
        ("7. A liquid sample has a mass of 27.0 g and a density of 2.70 g/mL. Calculate its volume.", "volume_2"),
        ("8. A liquid has a density of 0.789 g/mL and a volume of 25.0 mL. Calculate its mass.", "mass_3"),
        ("9. A student measures 2.58 g/mL. The accepted value is 2.70 g/mL. Calculate percent error.", "percent_error_4"),
    ]
    for prompt, answer_key in tasks:
        add_text(doc, prompt, 8.9, before=5, after=2)
        if key:
            add_answer(doc, DATA["answers"][answer_key])
        else:
            box = doc.add_table(rows=1, cols=1)
            set_table_widths(box, [7.18])
            set_cell_shading(box.cell(0, 0), WASH)
            clear_cell(box.cell(0, 0))
            add_cell_text(box.cell(0, 0), "Work space\n\n", 7.4, False, MUTED)
    add_text(doc, "10. Accuracy and precision", 11.2, True, NAVY, before=5, after=1)
    add_text(doc, "Accepted density: 2.70 g/mL", 8.8, after=2)
    data = doc.add_table(rows=2, cols=2)
    set_table_widths(data, [1.1, 6.08])
    values = [("Set A", "2.70, 2.71, 2.69 g/mL"), ("Set B", "2.42, 2.41, 2.42 g/mL")]
    for row, (label, vals) in zip(data.rows, values):
        set_cell_shading(row.cells[0], WASH)
        set_cell_shading(row.cells[1], WASH)
        clear_cell(row.cells[0]); clear_cell(row.cells[1])
        add_cell_text(row.cells[0], label, 8.3, True, NAVY)
        add_cell_text(row.cells[1], vals, 8.3)
    add_text(doc, "Which set is accurate? Which is precise? Explain using the pattern in the data.", 8.8, before=3)
    if key:
        add_answer(doc, DATA["answers"]["dataset"])
    else:
        add_lines(doc, 2)


def page4(doc, key=False):
    add_header(doc, 4, "Reason from Evidence", "Use the data to build a claim, diagnose the error, and plan a repair.", key)
    add_text(doc, "LAB CASE: UNKNOWN LIQUID", 11.5, True, NAVY, after=2)
    case = doc.add_table(rows=1, cols=1)
    set_table_widths(case, [7.18])
    set_cell_shading(case.cell(0, 0), WASH)
    clear_cell(case.cell(0, 0))
    add_cell_text(case.cell(0, 0), "A student uses an electronic balance and a graduated cylinder. The empty container has a mass of 42.31 g. The container plus liquid has a mass of 67.19 g. The liquid volume is 20.0 mL. The accepted density is 1.25 g/mL.", 8.8)
    q = doc.add_table(rows=2, cols=2)
    set_table_widths(q, [3.59, 3.59])
    prompts = [
        "11. Mass of liquid:" if key else "11. Mass of liquid: ____________________",
        "12. Experimental density:" if key else "12. Experimental density: ____________________",
        "13. Percent error:" if key else "13. Percent error: ____________________",
        "",
    ]
    for cell, prompt in zip([c for row in q.rows for c in row.cells], prompts):
        clear_cell(cell)
        if prompt:
            add_cell_text(cell, prompt, 8.3)
    if key:
        add_cell_text(q.cell(0, 0), DATA["answers"]["case_mass"], 8.2, True, ORANGE)
        add_cell_text(q.cell(0, 1), DATA["answers"]["case_density"], 8.2, True, ORANGE)
        add_cell_text(q.cell(1, 0), DATA["answers"]["case_percent_error"], 8.2, True, ORANGE)
    add_text(doc, "14. Claim - Evidence - Reasoning", 11.2, True, NAVY, before=4, after=2)
    for label, prompt in [
        ("Claim", "Is the experimental result close to the accepted value?"),
        ("Evidence", "Cite at least two numbers from the case."),
        ("Reasoning", "Explain what percent error tells you and what it does not tell you about precision."),
    ]:
        add_text(doc, f"{label}: {prompt}", 8.6, bold=(label == "Claim"), after=1)
        if not key:
            add_lines(doc, 1 if label != "Reasoning" else 2)
    if key:
        add_answer(doc, DATA["answers"]["case_claim"] + " Evidence: 1.244 g/mL before final rounding vs. 1.25 g/mL and 0.48% error.")
    add_text(doc, "15. Error diagnosis", 11.2, True, NAVY, before=4, after=2)
    add_text(doc, "Another student reads the liquid above eye level, rounds every intermediate value, and omits units. Name two problems and one exact repair action.", 8.6)
    if key:
        add_answer(doc, DATA["answers"]["repair"])
    else:
        add_lines(doc, 2)
    repair = doc.add_table(rows=2, cols=2)
    set_table_widths(repair, [3.59, 3.59])
    for row in repair.rows:
        for cell in row.cells:
            set_cell_shading(cell, NOTE)
            clear_cell(cell)
    labels = ["Item/skill to repair: __________________", "Next action: __________________________", "Incorrect idea or step: _______________", "When I will retry: ____________________"]
    for cell, label in zip([c for row in repair.rows for c in row.cells], labels):
        add_cell_text(cell, label, 7.8, False, NAVY)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.25)
    add_footer(section)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    for style_name in ("List Number", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(8.7)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    pages = [page1, page2, page3, page4]
    for key in (False, True):
        for idx, drawer in enumerate(pages):
            drawer(doc, key=key)
            if not (key and idx == len(pages) - 1):
                doc.add_page_break()
    core = doc.core_properties
    core.title = DATA["title"] + " - Editable"
    core.author = BRAND
    core.subject = "Chemistry measurement visual worksheet"
    core.keywords = "chemistry, measurement, lab tools, density, percent error"
    path = OUT / "CurioNest_Measurement_Visual_Worksheet_Editable.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    main()
